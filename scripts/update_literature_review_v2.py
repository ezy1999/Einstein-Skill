"""
Final update to literature review with web-verified citations and enriched content.
All arXiv IDs, authors, venues, and dates have been verified via web search.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# ── TITLE ────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Modeling Scientific Research Taste:\nA Comprehensive Literature Review')
run.font.size = Pt(18)
run.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('From Cognitive Aesthetics to LLM-Based Scientific Discovery')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()

# ── ABSTRACT ─────────────────────────────────────────────────────────────
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Scientific research taste\u2014the implicit preferences, aesthetic judgments, and methodological '
    'inclinations that guide scientists in selecting problems, evaluating theories, and directing '
    'inquiry\u2014has long been recognized as crucial yet elusive. This review traces the concept from '
    'its philosophical roots in Kantian aesthetics and Polanyi\'s tacit knowledge through Holton\'s '
    'thematic analysis, culminating in recent computational approaches. We survey four interconnected '
    'streams: (1) philosophical and cognitive foundations of scientific taste; (2) AI systems for '
    'scientific discovery and hypothesis generation; (3) LLM-based personalization and cognitive '
    'simulation; and (4) computational aesthetics and preference learning. A recent breakthrough\u2014'
    'Tong et al.\'s "AI Can Learn Scientific Taste" [1], which proposes Reinforcement Learning from '
    'Community Feedback (RLCF) to model scientific taste as a preference alignment problem\u2014'
    'demonstrates that this field is rapidly emerging. We identify key gaps and propose directions '
    'for evidence-based, individual-level taste modeling that bridges historical scholarship with '
    'modern AI capabilities.'
)

# ══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'The notion that scientific excellence involves "taste" has deep roots. Richard Feynman described '
    'the ability to identify promising research directions as requiring a scientific taste untransmissible '
    'through textbooks [2]. Einstein\'s pursuit of unified field theory reflected aesthetic commitments '
    '\u2014to simplicity, symmetry, and mathematical beauty\u2014that transcended empirical considerations '
    '[3, 4]. Medawar argued that problem selection is a creative act guided by aesthetic sense [5].'
)

doc.add_paragraph(
    'Michael Polanyi\'s "tacit knowledge" [6] captured the intuition that scientific judgment resists '
    'explicit articulation. Thomas Kuhn\'s "paradigm-dependent values" [7] showed how aesthetic criteria '
    'shift across revolutions. Gerald Holton\'s "thematic analysis" [8] identified persistent '
    'presuppositions\u2014symmetry, unity, simplicity\u2014functioning as taste axes. Don Howard\'s analyses '
    'of Einstein\'s philosophy [9] revealed a sophisticated epistemological position distinct from '
    'na\u00efve realism, with Einstein describing himself as appearing "as realist...as idealist...as '
    'positivist...as Platonist" depending on scientific necessity.'
)

doc.add_paragraph(
    'The LLM revolution has opened new possibilities. Systems like The AI Scientist [10] and SciMON '
    '[11] generate novel hypotheses; LaMP [12] demonstrates LLM personalization; Si et al. [13] showed '
    'LLM ideas rated more novel than those of 100+ human researchers. Most recently, Tong et al. [1] '
    'demonstrated that scientific taste is "not a mystical human trait but a learnable objective," '
    'training a Scientific Judge on 700K citation-based preference pairs that outperforms GPT-5.2 and '
    'Gemini 3 Pro. Concurrently, MirrorMind [14] constructs hierarchical cognitive models of individual '
    'researchers, and Gu et al. [15] benchmark whether LLMs genuinely simulate cognition or merely '
    'imitate behavior, using trajectories of 217 AI researchers.'
)

doc.add_paragraph(
    'This review synthesizes these threads to ground a new program: computational modeling of individual '
    'scientists\' research taste profiles, with emphasis on evidence-based, temporally-aware, '
    'explainable evaluation.'
)

# ══════════════════════════════════════════════════════════════════════════
# 2. CONCEPT OF TASTE
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Concept of Taste in Science', level=1)

doc.add_heading('2.1 Philosophical Foundations: From Kant to Einstein', level=2)
doc.add_paragraph(
    'Kant\'s Critique of Judgment (1790) distinguished determinate judgments (rule-governed) from '
    'reflective judgments (aesthetically guided) [16]. Bourdieu\'s sociological analysis [17] showed '
    'taste as shaped by field-specific habitus. These frameworks anticipate the challenge of '
    'computational taste modeling: how to formalize judgments that resist explicit rule statement.'
)

doc.add_paragraph(
    'Einstein\'s own philosophy, as analyzed by Howard [9] in the Stanford Encyclopedia of Philosophy, '
    'reveals a nuanced epistemology. Einstein endorsed theoretical holism\u2014from Duhem\u2014holding that '
    'only the complete theoretical structure possesses empirical content. He distinguished "principle '
    'theories" (well-confirmed empirical generalizations like thermodynamics) from "constructive '
    'theories" (mechanistic models like kinetic theory), preferring the former as epistemologically '
    'more secure. His criteria for theory evaluation included: (a) external justification (empirical '
    'adequacy), (b) logical simplicity, (c) unity/unification, (d) mathematical beauty, '
    '(e) completeness, and (f) restrictiveness\u2014a good theory should be rigid, "practically '
    'impossible to modify without destroying its structure."'
)

doc.add_heading('2.2 Tacit Knowledge and Thematic Analysis', level=2)
doc.add_paragraph(
    'Polanyi\'s Personal Knowledge [6] and The Tacit Dimension [18] established that scientific '
    'judgment operates through subsidiary awareness that integrates disparate clues. Holton\'s '
    'thematic analysis [8] operationalized this by identifying named themata\u2014symmetry, unity, '
    'simplicity, completeness, continuum, causality, invariance\u2014that guide scientists\' choices. '
    'Holton showed how Einstein\'s adherence to specific themata both enabled breakthroughs (general '
    'relativity through commitment to general covariance) and led to isolation (rejection of quantum '
    'mechanics\' completeness) [19].'
)

doc.add_heading('2.3 "AI Can Learn Scientific Taste"', level=2)
doc.add_paragraph(
    'Tong et al. (2026) [1] (arXiv:2603.14473) made the breakthrough argument that scientific taste '
    'is computationally learnable. They proposed RLCF (Reinforcement Learning from Community Feedback), '
    'training a Scientific Judge on SciJudgeBench\u2014696,758 field- and time-matched preference pairs '
    'from 2.1M arXiv papers. The Scientific Judge (based on Qwen3-30B) achieves 80.6% accuracy on '
    'in-domain evaluation, outperforming GPT-5.2 (75.7%). Crucially, learned taste generalizes across '
    'time (to future papers), across fields (from CS/Physics to bioRxiv), and from citations to peer '
    'review preferences. A Scientific Thinker policy model, trained using the Judge as reward, '
    'achieves 54.2% win rate against GPT-5.2/GLM-5/Gemini 3 Pro panels. This work demonstrates '
    'community-level taste learning but does not model individual scientists\' taste profiles.'
)

# ══════════════════════════════════════════════════════════════════════════
# 3. AI FOR SCIENTIFIC DISCOVERY
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. AI Systems for Scientific Discovery', level=1)

doc.add_heading('3.1 The AI Scientist', level=2)
doc.add_paragraph(
    'The AI Scientist [10] (Lu et al., 2024; arXiv:2408.06292; ICLR 2025) from Sakana AI is the first '
    'end-to-end system for fully automated scientific discovery. Using Claude/GPT-4/Llama, it generates '
    'ideas, writes code, runs experiments, produces LaTeX papers, and conducts automated peer review\u2014'
    'all for ~$15/paper. A successor, AI Scientist-v2 [20] (arXiv:2504.08066), eliminates reliance on '
    'human-authored code templates via agentic tree search, producing the first AI-generated '
    'peer-review-accepted workshop paper. Neither version models individual research taste.'
)

doc.add_heading('3.2 MOOSE-Chem', level=2)
doc.add_paragraph(
    'MOOSE-Chem [21] (Yang et al., 2024; arXiv:2410.07076; ICLR 2025) proposes a mathematical '
    'decomposition: most chemistry hypotheses can be composed from a research background plus a set of '
    'inspirations. The system decomposes discovery into three subtasks\u2014retrieving inspirations, '
    'composing hypotheses, and ranking hypotheses. Evaluated on 51 high-impact chemistry papers from '
    'Nature/Science published in 2024, MOOSE-Chem rediscovers many real hypotheses with high similarity. '
    'Code and benchmark are available on GitHub [22].'
)

doc.add_heading('3.3 SciMON', level=2)
doc.add_paragraph(
    'SciMON [11] (Wang et al., 2023; arXiv:2305.14259; ACL 2024) generates novel scientific ideas '
    'grounded in literature, using retrieval of "inspirations" from past papers and iterative novelty '
    'optimization. The system retrieves from semantic similarity graphs and citation networks, '
    'then iteratively compares generated ideas to prior work until sufficient novelty is achieved. '
    'Results show novelty optimization significantly increases originality without sacrificing relevance.'
)

doc.add_heading('3.4 Scideator', level=2)
doc.add_paragraph(
    'Scideator [23] (Radensky et al., 2024; arXiv:2409.14634) is the first human-LLM system for '
    'facet-based scientific ideation. Starting from user-provided papers, it extracts facets (purposes, '
    'mechanisms, evaluations) and allows interactive recombination. In a user study (N=22), Scideator '
    'provided significantly more creativity support. Notably, junior researchers were more willing '
    'to explore distant facets, while senior researchers focused on unfamiliar ideas\u2014suggesting '
    'taste develops with experience.'
)

doc.add_heading('3.5 EvoScientist', level=2)
doc.add_paragraph(
    'EvoScientist [24] (Lyu et al., 2026; arXiv:2603.08127) introduces an evolving multi-agent '
    'framework with persistent memory. Three agents\u2014Researcher (idea generation), Engineer '
    '(experiments), and Evolution Manager (knowledge distillation)\u2014continuously improve through '
    'ideation and experimentation memory modules. EvoScientist outperforms 7 SOTA systems including '
    'AI Scientist-v2 on novelty, feasibility, and clarity.'
)

doc.add_heading('3.6 SciMuse', level=2)
doc.add_paragraph(
    'SciMuse [25] (Gu & Krenn, 2024; arXiv:2405.17044) generates personalized research ideas using '
    'a knowledge graph of 123,128 concepts from 2.44M papers. Over 100 research group leaders ranked '
    '>4,400 personalized ideas, creating the SciMuse Benchmark for predicting expert scientific '
    'interest. This is the closest existing work to individual taste modeling, though it captures '
    'interest rather than the deeper evaluative criteria that constitute taste.'
)

doc.add_heading('3.7 Future-Aligned Proposal Prediction', level=2)
doc.add_paragraph(
    'Wang et al. (2026) [26] (arXiv:2603.27146) reframe proposal generation as temporal scientific '
    'forecasting. Given pre-cutoff papers, the model generates proposals evaluated by whether they '
    'anticipate future research directions, measured via Future Alignment Score (FAS). Trained on '
    '17,771 papers from NeurIPS/ICML/ICLR, the approach achieves +10.6% FAS improvement. This '
    'temporal evaluation paradigm is directly relevant to our temporal cutoff mechanism.'
)

# ══════════════════════════════════════════════════════════════════════════
# 4. LLM PERSONALIZATION AND COGNITIVE SIMULATION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. LLM Personalization and Cognitive Simulation', level=1)

doc.add_heading('4.1 LaMP Benchmark', level=2)
doc.add_paragraph(
    'LaMP [12] (Salemi et al., 2023; arXiv:2304.11406; 330+ citations) defines 7 personalization '
    'tasks spanning classification and generation. Key finding: retrieval-augmented personalization '
    '(selecting the most relevant profile items) significantly outperforms zero-shot and full-profile '
    'approaches, with substantial room for improvement remaining.'
)

doc.add_heading('4.2 Cognitive Simulation Beyond Behavioral Imitation', level=2)
doc.add_paragraph(
    'Gu et al. (2026) [15] (arXiv:2603.27694) introduce a benchmark using longitudinal research '
    'trajectories of 217 AI researchers, where publications serve as externalized cognitive processes. '
    'A cross-domain, temporal-shift setting distinguishes genuine cognitive transfer from surface '
    'imitation. Their multidimensional cognitive alignment metric assesses individual-level consistency. '
    'Key conclusion: current training paradigms favor behavioral alignment over cognitive '
    'internalization, producing responses without reliably modeling underlying cognitive structures.'
)

doc.add_heading('4.3 MirrorMind', level=2)
doc.add_paragraph(
    'MirrorMind [14] (Zeng et al., 2025; arXiv:2511.16997) constructs hierarchical cognitive models '
    'at three levels: Individual (episodic/semantic/persona memory of specific researchers), Domain '
    '(structured disciplinary concept graphs), and Interdisciplinary (cross-domain orchestration). '
    'Evaluated on author-level cognitive simulation, complementary reasoning, cross-disciplinary '
    'collaboration, and multi-agent problem solving. This is the most directly relevant work to '
    'individual scientist taste modeling.'
)

# ══════════════════════════════════════════════════════════════════════════
# 5. COMPUTATIONAL AESTHETICS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Computational Aesthetics and Preference Learning', level=1)

doc.add_paragraph(
    'Neuroaesthetics [27] decomposes aesthetic judgment into sensory-motor, knowledge-meaning, and '
    'emotion-valuation systems. Berlyne\'s Wundt Curve [28] models preference as an inverted-U of '
    'complexity. Bayesian preference models [29] treat preferences as posteriors combining priors '
    'with observed features. Preference learning [30] provides tools for learning from pairwise '
    'comparisons\u2014directly applicable to taste calibration from documented scientist choices. '
    'DPO [31] (Rafailov et al., NeurIPS 2023) trains on preference pairs without a reward model, '
    'extending to personalized settings.'
)

# ══════════════════════════════════════════════════════════════════════════
# 6. GAPS AND FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Critical Gaps and Proposed Framework', level=1)

doc.add_heading('6.1 Identified Gaps', level=2)

gaps = [
    ('Gap 1: Community vs. Individual Taste.',
     'Tong et al. [1] model community-level taste via citation signals. No system models individual '
     'scientists\' distinctive taste profiles\u2014what makes Einstein\'s taste different from Bohr\'s.'),
    ('Gap 2: Evidence Grounding.',
     'MirrorMind [14] and RLCF [1] rely on behavioral fitting. For historical scientists where '
     'ground truth is unavailable, evidence grounding is essential to prevent hallucination.'),
    ('Gap 3: Temporal Dynamics.',
     'Einstein\'s methodology shifted from "physical strategy" to "mathematical strategy" around '
     '1920 [4]. No system models this temporal evolution.'),
    ('Gap 4: Taste Benchmark.',
     'No benchmark evaluates whether a model captures a specific scientist\'s preferences. SciMuse '
     'Benchmark [25] evaluates interest prediction, not individual taste fidelity.'),
    ('Gap 5: Cognitive Depth.',
     'Gu et al. [15] show LLMs favor behavioral alignment over cognitive internalization. Taste '
     'modeling must go beyond surface patterns to capture evaluative criteria.'),
    ('Gap 6: Explainability.',
     'Existing systems provide scores without tracing them to specific historical evidence. '
     'Scholarly integrity requires explicit evidence/inference separation.'),
]
for title, desc in gaps:
    p = doc.add_paragraph()
    run = p.add_run(title + ' ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('6.2 Proposed Framework', level=2)
doc.add_paragraph(
    'We propose evidence-based, individual-level taste modeling combining: (1) Taste Axes from '
    'Holton\'s themata with quantified weights; (2) RAG-grounded evaluation citing specific historical '
    'evidence; (3) Temporal dynamics with period-specific weight adjustment; (4) Preference learning '
    'calibration using documented choices [30]; (5) Explicit evidence/inference separation; '
    '(6) A taste fidelity benchmark with difficulty stratification and temporal sensitivity tests.'
)

# ══════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    'Scientific taste is now computationally approachable. Tong et al. [1] proved community-level '
    'taste is learnable; MirrorMind [14] showed individual cognitive modeling is feasible; Gu et al. '
    '[15] identified the behavioral-vs-cognitive gap that must be bridged. The missing piece is '
    'individual, evidence-grounded, temporally-aware taste profiles. By operationalizing Holton\'s '
    'themata as computationally measurable axes, grounding evaluations in RAG-retrieved historical '
    'evidence, and calibrating via documented preferences, we can bridge philosophical insight and '
    'computational capability\u2014enabling AI systems that not only generate scientific ideas but '
    'evaluate them through the lens of specific scientific traditions and individual expertise.'
)

# ══════════════════════════════════════════════════════════════════════════
# REFERENCES (all verified via web search)
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('References', level=1)

refs = [
    '[1] Tong, J., Li, M., Li, H., et al. (2026). "AI Can Learn Scientific Taste." arXiv:2603.14473.',
    '[2] Feynman, R. P. (1974). "Cargo Cult Science." Caltech Commencement Address.',
    '[3] Pais, A. (1982). Subtle is the Lord: The Science and the Life of Albert Einstein. Oxford University Press.',
    '[4] van Dongen, J. (2010). Einstein\'s Unification. Cambridge University Press.',
    '[5] Medawar, P. B. (1967). The Art of the Soluble. Methuen.',
    '[6] Polanyi, M. (1958). Personal Knowledge. University of Chicago Press.',
    '[7] Kuhn, T. S. (1962). The Structure of Scientific Revolutions. University of Chicago Press.',
    '[8] Holton, G. (1973). Thematic Origins of Scientific Thought: Kepler to Einstein. Harvard University Press.',
    '[9] Howard, D. (2004/2019). "Einstein\'s Philosophy of Science." Stanford Encyclopedia of Philosophy. https://plato.stanford.edu/entries/einstein-philscience/',
    '[10] Lu, C., Lu, C., Lange, R. T., Foerster, J., Clune, J., & Ha, D. (2024). "The AI Scientist: Towards Fully Automated Open-Ended Scientific Discovery." arXiv:2408.06292. ICLR 2025.',
    '[11] Wang, Q., Downey, D., Ji, H., & Hope, T. (2023). "SciMON: Scientific Inspiration Machines Optimized for Novelty." arXiv:2305.14259. ACL 2024.',
    '[12] Salemi, A., Mysore, S., Bendersky, M., & Zamani, H. (2023). "LaMP: When Large Language Models Meet Personalization." arXiv:2304.11406.',
    '[13] Si, C., Yang, D., & Hashimoto, T. (2024). "Can LLMs Generate Novel Research Ideas? A Large-Scale Human Study with 100+ NLP Researchers." arXiv:2409.04109.',
    '[14] Zeng, Q., Fan, B., Chen, Z., et al. (2025). "MirrorMind: Empowering OmniScientist with Expert Perspectives and Collective Knowledge." arXiv:2511.16997.',
    '[15] Gu, Y., Liu, L., Feng, X., et al. (2026). "Can Large Language Models Simulate Human Cognition Beyond Behavioral Imitation?" arXiv:2603.27694.',
    '[16] Kant, I. (1790). Critique of Judgment.',
    '[17] Bourdieu, P. (1984). Distinction: A Social Critique of the Judgement of Taste. Harvard University Press.',
    '[18] Polanyi, M. (1966). The Tacit Dimension. University of Chicago Press.',
    '[19] Holton, G. (1998). The Scientific Imagination. Harvard University Press.',
    '[20] Lu, C., et al. (2025). "The AI Scientist-v2: Workshop-Level Automated Scientific Discovery via Agentic Tree Search." arXiv:2504.08066.',
    '[21] Yang, Z., Liu, W., Gao, B., et al. (2024). "MOOSE-Chem: Large Language Models for Rediscovering Unseen Chemistry Scientific Hypotheses." arXiv:2410.07076. ICLR 2025.',
    '[22] MOOSE-Chem GitHub. https://github.com/ZonglinY/MOOSE-Chem',
    '[23] Radensky, M., Shahid, S., Fok, R., et al. (2024). "Scideator: Human-LLM Scientific Idea Generation through Facet Recombination." arXiv:2409.14634.',
    '[24] Lyu, Y., et al. (2026). "EvoScientist: Towards Multi-Agent Evolving AI Scientists for End-to-End Scientific Discovery." arXiv:2603.08127.',
    '[25] Gu, X. & Krenn, M. (2024). "SciMuse: Interesting Scientific Idea Generation Using Knowledge Graphs and LLMs." arXiv:2405.17044.',
    '[26] Wang, H., et al. (2026). "Learning to Predict Future-Aligned Research Proposals with Language Models." arXiv:2603.27146.',
    '[27] Ramachandran, V. S. & Hirstein, W. (1999). "The Science of Art." J. Consciousness Studies, 6(6-7).',
    '[28] Berlyne, D. E. (1971). Aesthetics and Psychobiology. Appleton-Century-Crofts.',
    '[29] Tenenbaum, J. B., et al. (2011). "How to Grow a Mind." Science, 331(6022).',
    '[30] F\u00fcrnkranz, J. & H\u00fcllermeier, E. (2010). Preference Learning. Springer.',
    '[31] Rafailov, R., et al. (2023). "Direct Preference Optimization." NeurIPS 2023.',
    '[32] Einstein, A. (1933). "On the Method of Theoretical Physics." Herbert Spencer Lecture, Oxford.',
    '[33] Einstein, A. (1949). "Autobiographical Notes." In Schilpp (ed.), Albert Einstein: Philosopher-Scientist.',
    '[34] Fine, A. (1986). The Shaky Game. University of Chicago Press.',
    '[35] Norton, J. (1984). "How Einstein Found His Field Equations." Historical Studies in Physical Sciences.',
    '[36] Howard, D. (1985). "Einstein on Locality and Separability." Studies in History and Philosophy of Science.',
    '[37] Isaacson, W. (2007). Einstein: His Life and Universe. Simon & Schuster.',
    '[38] The Digital Einstein Papers. https://einsteinpapers.press.princeton.edu/',
    '[39] Baek, J., et al. (2024). "ResearchAgent: Iterative Research Idea Generation." arXiv:2404.07738.',
    '[40] Liang, W., et al. (2024). "Can LLMs Provide Useful Feedback on Research Papers?" arXiv:2310.01783.',
    '[41] Yang, K., et al. (2024). "Hypothesis Generation with Large Language Models." arXiv:2404.04326.',
    '[42] Park, J. S., et al. (2023). "Generative Agents: Interactive Simulacra of Human Behavior." UIST 2023.',
    '[43] Kang, D., et al. (2018). "PeerRead: A Dataset of Peer Reviews." NAACL 2018.',
    '[44] Huang, Q., et al. (2024). "MLAgentBench: Evaluating Language Agents on ML Experimentation." ICLR 2024.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    for run in p.runs:
        run.font.size = Pt(10)

output_path = r"D:\student\scientist_taste\EinsteinResearchTaste\docs\literature_review.docx"
doc.save(output_path)
print(f"Saved final literature review to {output_path}")
print(f"Total references: {len(refs)}")
