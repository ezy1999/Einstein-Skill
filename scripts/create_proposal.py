"""
Generate the Project Proposal Word document.
Covers data collection, taste modeling, benchmark construction, and implementation details.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Research Taste Modeling: Project Proposal\nData Collection, Modeling, Benchmark, and Implementation Plan')
run.font.size = Pt(18)
run.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)

doc.add_paragraph(
    'This document presents the complete technical plan for building a computational system '
    'that models the research taste of historical scientists, with Albert Einstein as the '
    'initial case study. The system is designed to: (1) evaluate and rank candidate scientific '
    'theories/questions according to a scientist\'s documented preferences; (2) provide '
    'evidence-grounded explanations for each evaluation; (3) strictly separate historical '
    'evidence from model inference; and (4) support temporal cutoffs to prevent anachronism.'
)

doc.add_paragraph(
    'The project proceeds in four phases: data collection, mathematical formalization, '
    'benchmark construction, and system implementation. Each phase is described in detail below.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: DATA COLLECTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Data Collection Plan', level=1)

doc.add_heading('2.1 Data Sources', level=2)
doc.add_paragraph(
    'We categorize data sources by their proximity to the scientist\'s actual thinking:'
)

# Table: Data sources
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Source Type', 'Examples', 'Priority', 'Expected Volume']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['Primary (Own writings)', 'Published papers, lectures, letters, notebooks, autobiographical notes',
     'Critical', '~500 documents'],
    ['Secondary (Scholarly)', 'Pais (1982), Holton (1973), van Dongen (2010), Norton (1984), Howard (1985), Fine (1986)',
     'High', '~50 books/papers'],
    ['Tertiary (General)', 'Isaacson (2007), popular biographies, textbook accounts',
     'Medium', '~20 sources'],
    ['Contextual', 'Contemporary scientific papers Einstein engaged with, Solvay proceedings, peer reviews',
     'Medium', '~200 documents'],
]
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        table.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('2.2 Specific Data Sources for Einstein', level=2)

doc.add_paragraph(
    'Primary sources (Einstein\'s own words):'
)
items = [
    'The Collected Papers of Albert Einstein (Princeton University Press, 15+ volumes) \u2014 '
    'Full text available at einsteinpapers.press.princeton.edu',
    'Einstein Archives Online (Hebrew University) \u2014 Metadata and selected digitized manuscripts',
    'Einstein-Born Letters (1916\u20131955) \u2014 Key source for quantum mechanics debates',
    'Einstein-Besso Correspondence \u2014 Candid discussions of scientific methodology',
    '"Autobiographical Notes" in Schilpp (1949) \u2014 Einstein\'s own account of his intellectual development',
    '"On the Method of Theoretical Physics" (1933) \u2014 Einstein\'s most explicit methodological statement',
    '"Physics and Reality" (1936) \u2014 Einstein\'s philosophy of physical theory',
    '"Geometry and Experience" (1921) \u2014 On the relationship between mathematics and reality',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Secondary sources (scholarly analyses of Einstein\'s taste):'
)
items2 = [
    'Holton, G. (1973) Thematic Origins of Scientific Thought \u2014 Identifies Einstein\'s core themata',
    'Pais, A. (1982) Subtle is the Lord \u2014 Definitive scientific biography',
    'van Dongen, J. (2010) Einstein\'s Unification \u2014 Documents the "physical strategy" to "mathematical strategy" shift',
    'Norton, J. (1984, 1991, 2000) \u2014 Multiple papers on Einstein\'s methodology and thought experiments',
    'Howard, D. (1985, 1990) \u2014 Einstein on locality, separability, and quantum theory',
    'Fine, A. (1986) The Shaky Game \u2014 Einstein\'s realism and quantum debates',
    'Stachel, J. (2002) Einstein from \'B\' to \'Z\' \u2014 Scholarly essays on Einstein\'s development',
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 Data Processing Pipeline', level=2)
doc.add_paragraph(
    'Raw data is processed into structured "evidence records" with the following fields:'
)

table2 = doc.add_table(rows=10, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Field', 'Type', 'Description']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

fields = [
    ['id', 'string', 'Unique identifier'],
    ['content', 'text', 'The evidence text/observation'],
    ['source_text', 'string', 'Full citation'],
    ['source_type', 'enum', 'primary|secondary|tertiary|inference'],
    ['confidence', 'enum', 'direct|strong|moderate|weak|speculative'],
    ['year', 'int?', 'Year the evidence pertains to'],
    ['period', 'string?', 'Named career period'],
    ['relevant_axes', 'list[str]', 'Taste axes this evidence informs'],
    ['axis_valence', 'dict', 'Per-axis support/contradiction scores'],
]
for i, row_data in enumerate(fields):
    for j, val in enumerate(row_data):
        table2.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'The processing pipeline: (1) Extract text/quotes from sources \u2192 (2) Annotate with '
    'source type and confidence level \u2192 (3) Tag relevant taste axes \u2192 (4) Assign temporal '
    'period \u2192 (5) Compute axis valence scores \u2192 (6) Generate embeddings for retrieval.'
)

doc.add_heading('2.4 Data Quality Controls', level=2)
doc.add_paragraph(
    'Each evidence record undergoes quality control: (a) Source verification \u2014 primary quotes '
    'are cross-referenced against original publications; (b) Confidence calibration \u2014 two '
    'independent annotators rate confidence level; (c) Axis relevance validation \u2014 tagged axes '
    'must be justified by the content; (d) Temporal accuracy \u2014 years are verified against '
    'publication records and chronologies.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: MATHEMATICAL FORMALIZATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Mathematical Formalization of Research Taste', level=1)

doc.add_heading('3.1 Taste as a Vector in Axis Space', level=2)
doc.add_paragraph(
    'We define a scientist\'s research taste T at time t as a vector in an n-dimensional '
    'taste axis space:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('T(t) = [w\u2081(t)\u00b7a\u2081, w\u2082(t)\u00b7a\u2082, ..., w\u2099(t)\u00b7a\u2099]')
run.font.name = 'Cambria Math'
run.italic = True

doc.add_paragraph(
    'where a\u1d62 represents the i-th taste axis (e.g., simplicity, unity, invariance) and '
    'w\u1d62(t) represents the time-dependent weight of that axis. The axes themselves are defined '
    'by their associated evidence corpus and keyword descriptors.'
)

doc.add_heading('3.2 Candidate Evaluation Function', level=2)
doc.add_paragraph(
    'Given a candidate theory/question c and a taste profile T(t), the evaluation function '
    'computes an alignment score:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('S(c, T, t) = \u03a3\u1d62 w\u1d62(t) \u00b7 f(c, a\u1d62, E\u1d62(t))')
run.font.name = 'Cambria Math'
run.italic = True

doc.add_paragraph(
    'where f(c, a\u1d62, E\u1d62(t)) is the axis-specific scoring function that measures how well '
    'candidate c aligns with axis a\u1d62 given evidence E\u1d62(t) available up to time t. '
    'The scoring function f is implemented via LLM-based evaluation grounded in retrieved evidence.'
)

doc.add_heading('3.3 Temporal Weight Evolution', level=2)
doc.add_paragraph(
    'Axis weights evolve over career periods. We model this as a piecewise function with smooth '
    'transitions between periods. For Einstein, we identify four major periods with documented '
    'shifts in methodological emphasis:'
)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'
headers3 = ['Period', 'Years', 'Dominant Axes']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

periods = [
    ['Early Revolutionary', '1900\u20131905', 'Empirical grounding, Thought experiment, Simplicity'],
    ['General Relativity', '1906\u20131915', 'Invariance, Unity, Causal continuity'],
    ['Quantum Debates', '1916\u20131935', 'Physical reality, Causal continuity, Mathematical beauty'],
    ['Unified Field Theory', '1936\u20131955', 'Unity, Mathematical beauty, Simplicity'],
]
for i, row_data in enumerate(periods):
    for j, val in enumerate(row_data):
        table3.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('3.4 Evidence-Grounded Scoring', level=2)
doc.add_paragraph(
    'The axis scoring function f decomposes into: (a) an evidence retrieval step that finds '
    'relevant historical evidence using semantic similarity; (b) an LLM evaluation step that '
    'scores the candidate against the axis given the retrieved evidence; (c) a confidence '
    'estimation step that reflects the quality and quantity of supporting evidence. Crucially, '
    'scores without supporting evidence are flagged as "model inference" rather than '
    '"evidence-based," maintaining the distinction required for scholarly integrity.'
)

doc.add_heading('3.5 Fitting the Taste Profile', level=2)
doc.add_paragraph(
    'The taste profile can be "fitted" to a scientist using a two-stage process:'
)
doc.add_paragraph(
    'Stage 1 (Evidence-based initialization): Extract initial axis weights from the relative '
    'frequency and strength of evidence for each axis in each period. This gives a prior '
    'distribution over weights based purely on documented evidence.',
)
doc.add_paragraph(
    'Stage 2 (Behavioral calibration): Refine weights using known preference data \u2014 '
    'documented cases where the scientist chose between alternatives (e.g., Einstein choosing '
    'GR over Nordstr\u00f6m\'s scalar theory). Optimize weights to reproduce known choices via '
    'pairwise ranking loss:',
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('L = \u03a3\u2c7c max(0, S(c\u207b, T, t) - S(c\u207a, T, t) + \u03b5)')
run.font.name = 'Cambria Math'
run.italic = True

doc.add_paragraph(
    'where c\u207a is the preferred candidate and c\u207b is the rejected one in documented '
    'historical choices, and \u03b5 is a margin parameter.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: BENCHMARK CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Benchmark Construction', level=1)

doc.add_heading('4.1 Benchmark Design Principles', level=2)
doc.add_paragraph(
    'The benchmark evaluates whether the taste model accurately captures a scientist\'s '
    'documented preferences. Design principles:'
)

items3 = [
    'Historical grounding: Every benchmark case must be supported by documented evidence '
    'of the scientist\'s actual preference.',
    'Difficulty stratification: Cases are organized by difficulty (easy/medium/hard) based '
    'on the ambiguity of the historical record.',
    'Temporal sensitivity: Some cases test whether the model correctly adjusts for different '
    'career periods.',
    'Axis coverage: The benchmark covers all defined taste axes, with multiple cases per axis.',
    'Negative testing: Include cases where the scientist\'s preference contradicts the naive '
    'expectation (e.g., Einstein sometimes prioritizing mathematical beauty over simplicity).',
]
for item in items3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 Benchmark Case Types', level=2)

table4 = doc.add_table(rows=5, cols=4)
table4.style = 'Light Grid Accent 1'
headers4 = ['Type', 'Description', 'Metric', 'Example']
for i, h in enumerate(headers4):
    table4.rows[0].cells[i].text = h
    for p in table4.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

cases = [
    ['Binary Preference', 'A vs B, known preference', 'Accuracy', 'SR vs Lorentz ether theory'],
    ['Ranking', 'Rank N candidates, known order', 'Kendall \u03c4', 'Rank interpretations of QM'],
    ['Temporal', 'Same question, different cutoffs', 'Score difference sign', 'Mathematical beauty pre/post 1920'],
    ['Axis Isolation', 'Test single axis scoring', 'Correlation with expert rating', 'Unity score for unification attempts'],
]
for i, row_data in enumerate(cases):
    for j, val in enumerate(row_data):
        table4.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('4.3 Sample Selection', level=2)
doc.add_paragraph(
    'Benchmark samples are selected as follows:'
)
doc.add_paragraph(
    'Training set (for weight fitting): 20 documented preference pairs, balanced across '
    'taste axes and career periods. Sources: Einstein\'s explicit methodological statements, '
    'documented theory choices (e.g., choosing GR over competitors), and expressed evaluations '
    'of others\' work (e.g., responses to quantum mechanics).',
)
doc.add_paragraph(
    'Validation set: 10 preference pairs used for hyperparameter tuning (number of evidence '
    'records, LLM temperature, confidence thresholds).',
)
doc.add_paragraph(
    'Test set: 15 held-out preference pairs with difficulty stratification: 5 easy (clear '
    'historical preferences), 5 medium (nuanced/context-dependent), 5 hard (temporal sensitivity '
    'or contradictory evidence).',
)
doc.add_paragraph(
    'Data split criteria: (a) No temporal leakage \u2014 training cases from earlier periods than '
    'test cases where possible; (b) No evidence overlap \u2014 evidence used to construct training '
    'cases is excluded from test case retrieval; (c) Balanced axis coverage \u2014 each axis appears '
    'in at least 3 cases per split.',
)

doc.add_heading('4.4 Evaluation Metrics', level=2)
doc.add_paragraph(
    'Primary metrics: (1) Pairwise accuracy \u2014 fraction of binary preferences correctly predicted; '
    '(2) Kendall \u03c4 rank correlation \u2014 for ranking tasks; (3) Evidence utilization \u2014 '
    'fraction of scores grounded in retrieved evidence. Secondary metrics: (4) Calibration \u2014 '
    'correlation between model confidence and accuracy; (5) Temporal sensitivity \u2014 score '
    'change in expected direction across periods; (6) Axis decomposability \u2014 do individual '
    'axis scores contribute meaningful information beyond the aggregate?'
)

doc.add_heading('4.5 Ablation Studies', level=2)
doc.add_paragraph('Planned ablation experiments:')
ablations = [
    'No evidence retrieval (LLM only): Test whether RAG improves over pure LLM reasoning.',
    'No temporal cutoff: Test whether temporal awareness improves accuracy.',
    'Equal weights (no period adjustment): Test whether period-specific weights matter.',
    'Single axis only: Test each axis in isolation for independent contribution.',
    'Primary sources only vs. all sources: Test whether secondary sources add value.',
    'Embedding retrieval vs. keyword retrieval: Compare retrieval strategies.',
    'Different LLM backbones: Compare Claude, GPT-4, Llama across metrics.',
    'Evidence quantity: Vary top-k from 1 to 20 to find optimal evidence density.',
]
for item in ablations:
    doc.add_paragraph(item, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: IMPLEMENTATION PLAN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. System Design and Implementation', level=1)

doc.add_heading('5.1 Architecture Overview', level=2)
doc.add_paragraph(
    'The system follows a Retrieval-Augmented Generation (RAG) architecture with these components:'
)

items4 = [
    'Evidence Store: Structured database of historical evidence records with metadata',
    'Evidence Retriever: Semantic search over the evidence corpus (sentence-transformers + FAISS)',
    'Taste Evaluator: LLM-based scoring of candidates against taste axes',
    'Taste Scorer: Deterministic aggregation of axis scores with period-specific weights',
    'Pipeline Orchestrator: End-to-end coordination of retrieval, evaluation, and scoring',
    'Agent Interface: Conversational wrapper for interactive use',
    'Skill Wrapper: Claude Code skill for community distribution',
]
for item in items4:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 LLM Integration Strategy', level=2)
doc.add_paragraph(
    'The system uses LLMs in a carefully constrained manner:'
)
doc.add_paragraph(
    '(1) System prompt engineering: The LLM receives a detailed system prompt that: '
    '(a) defines all taste axes with historical grounding; (b) enforces temporal cutoff awareness; '
    '(c) requires evidence citation for each score; (d) explicitly prohibits role-playing; '
    '(e) mandates separate reporting of evidence-based vs. inferred assessments.',
)
doc.add_paragraph(
    '(2) Structured output: The LLM produces JSON-formatted responses with per-axis scores, '
    'evidence IDs, confidence levels, and explanations. This ensures machine-parseable output '
    'while maintaining interpretability.',
)
doc.add_paragraph(
    '(3) Evidence grounding: The LLM receives retrieved evidence as context, anchoring its '
    'evaluations in specific historical records rather than relying on parametric knowledge.',
)

doc.add_heading('5.3 Key Implementation Decisions', level=2)

table5 = doc.add_table(rows=8, cols=3)
table5.style = 'Light Grid Accent 1'
headers5 = ['Decision', 'Choice', 'Rationale']
for i, h in enumerate(headers5):
    table5.rows[0].cells[i].text = h
    for p in table5.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

decisions = [
    ['LLM Provider', 'Anthropic Claude (primary), OpenAI GPT-4 (secondary)',
     'Claude\'s strong instruction following and JSON output'],
    ['Embedding Model', 'sentence-transformers (all-MiniLM-L6-v2)',
     'Good balance of quality and speed for evidence retrieval'],
    ['Vector Store', 'FAISS (CPU)', 'No GPU required, sufficient for <10K records'],
    ['Scoring', 'Weighted axis aggregation', 'Interpretable, evidence-traceable'],
    ['Temporal Model', 'Piecewise constant with period boundaries',
     'Matches documented career periods'],
    ['Output Format', 'Structured JSON with evidence links',
     'Machine-parseable, auditable'],
    ['Data Format', 'JSON/JSONL evidence records',
     'Simple, portable, version-controllable'],
]
for i, row_data in enumerate(decisions):
    for j, val in enumerate(row_data):
        table5.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('5.4 Project Module Structure', level=2)
doc.add_paragraph(
    'The Python package (einstein_taste) is organized as follows:'
)

modules = [
    'config/settings.py \u2014 Taste axes definitions, temporal periods, model configuration',
    'core/evidence.py \u2014 Evidence record data model and store with filtering',
    'core/retriever.py \u2014 RAG retrieval component (keyword + semantic search)',
    'core/taste_model.py \u2014 Taste scoring model with axis aggregation',
    'core/evaluator.py \u2014 LLM-based axis evaluation with prompt engineering',
    'core/pipeline.py \u2014 End-to-end orchestration pipeline',
    'data/loader.py \u2014 Data loading and serialization',
    'data/fetcher.py \u2014 Online data fetching from public sources',
    'data/seed_evidence.py \u2014 Built-in seed evidence corpus',
    'evaluation/benchmark.py \u2014 Benchmark cases and evaluation metrics',
    'agents/taste_agent.py \u2014 Conversational agent wrapper',
    'skills/taste_skill.py \u2014 Claude Code skill interface',
    'utils/formatting.py \u2014 Output formatting utilities',
    'cli.py \u2014 Command-line interface',
]
for item in modules:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.5 Extensibility: Adding New Scientists', level=2)
doc.add_paragraph(
    'The system is designed for extensibility. To add a new scientist (e.g., Feynman, Newton, '
    'Mendel), one needs to: (1) Define taste axes specific to that scientist (may overlap with '
    'Einstein\'s axes but with different weights); (2) Collect evidence records from their '
    'writings and scholarly analyses; (3) Define temporal periods with dominant axes; '
    '(4) Create benchmark cases from documented preferences; (5) Calibrate axis weights '
    'using the pairwise ranking loss.'
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Development Roadmap', level=1)

table6 = doc.add_table(rows=5, cols=3)
table6.style = 'Light Grid Accent 1'
headers6 = ['Phase', 'Timeline', 'Deliverables']
for i, h in enumerate(headers6):
    table6.rows[0].cells[i].text = h
    for p in table6.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

phases = [
    ['Phase 1: Prototype', 'Weeks 1\u20132',
     'Working pipeline with seed data, CLI, basic benchmark (6 cases)'],
    ['Phase 2: Data Enrichment', 'Weeks 3\u20134',
     'Full evidence corpus (200+ records), semantic retrieval, expanded benchmark (45 cases)'],
    ['Phase 3: Calibration', 'Weeks 5\u20136',
     'Weight fitting via pairwise loss, ablation experiments, multi-LLM comparison'],
    ['Phase 4: Extension', 'Weeks 7\u20138',
     'Add second scientist (Feynman), cross-scientist comparison, community skill release'],
]
for i, row_data in enumerate(phases):
    for j, val in enumerate(row_data):
        table6.rows[i+1].cells[j].text = val

doc.add_paragraph()

# ── Save ────────────────────────────────────────────────────────────────────
output_dir = r"D:\student\scientist_taste\EinsteinResearchTaste\docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "project_proposal.docx")
doc.save(output_path)
print(f"Saved project proposal to {output_path}")
