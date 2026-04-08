"""
Generate the Literature Review Word document.
Uses python-docx to create a comprehensive review in journal introduction format.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 12, True)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Modeling Scientific Research Taste:\nA Comprehensive Literature Review')
run.font.size = Pt(18)
run.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('From Cognitive Aesthetics to LLM-Based Scientific Discovery')
run.font.size = Pt(14)
run.italic = True
run.font.name = 'Times New Roman'

doc.add_paragraph()  # spacer

# ── Abstract ────────────────────────────────────────────────────────────────
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Scientific research taste\u2014the implicit preferences, aesthetic judgments, and methodological '
    'inclinations that guide scientists in selecting problems, evaluating theories, and directing '
    'inquiry\u2014has long been recognized as a crucial yet elusive component of scientific excellence. '
    'This review traces the concept of taste from its philosophical roots in Kantian aesthetics and '
    'Polanyi\'s tacit knowledge through modern cognitive science frameworks, culminating in recent '
    'computational approaches using large language models (LLMs). We survey three interconnected '
    'research streams: (1) the philosophical and cognitive science foundations of scientific taste; '
    '(2) AI systems for autonomous scientific discovery and hypothesis generation; and '
    '(3) LLM-based personalization and cognitive simulation. We identify key gaps in the literature, '
    'particularly the absence of formal computational models that capture individual scientists\' '
    'research taste profiles, and propose directions for evidence-based taste modeling that bridges '
    'historical scholarship with modern AI capabilities.'
)

# ── 1. Introduction ─────────────────────────────────────────────────────────
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'The notion that scientific excellence involves a form of "taste" has deep roots in the history '
    'and philosophy of science. Richard Feynman famously described the ability to identify '
    'promising research directions as requiring a kind of scientific taste that cannot be taught '
    'through textbooks alone [1]. Albert Einstein\'s persistent pursuit of unified field theory, '
    'despite decades of failure, reflected aesthetic commitments\u2014to simplicity, symmetry, and '
    'mathematical beauty\u2014that transcended purely empirical considerations [2, 3]. Peter Medawar '
    'argued that the selection of a scientific problem is itself a creative act guided by an '
    '"aesthetic sense" for what is important and tractable [4].'
)

doc.add_paragraph(
    'Despite its acknowledged importance, scientific taste has resisted formal characterization. '
    'Michael Polanyi\'s concept of "tacit knowledge" [5] captured the intuition that much of what '
    'guides expert scientific judgment is difficult to articulate explicitly. Thomas Kuhn\'s notion '
    'of "paradigm-dependent values" [6] highlighted how aesthetic criteria shift across scientific '
    'revolutions. Gerald Holton\'s "thematic analysis" [7] identified persistent presuppositions '
    '(themata) that constrain scientists\' theoretical choices\u2014symmetry, unity, simplicity\u2014which '
    'function as taste axes guiding inquiry.'
)

doc.add_paragraph(
    'The recent revolution in large language models (LLMs) has opened new possibilities for '
    'computational approaches to scientific taste. Systems like The AI Scientist [8] and '
    'SciMON [9] demonstrate that LLMs can generate novel scientific hypotheses, while '
    'personalization benchmarks like LaMP [10] show that LLMs can be adapted to capture individual '
    'preferences. Concurrently, work on LLM-based cognitive simulation [11, 12] raises the '
    'question of whether these models can go beyond behavioral mimicry to capture deeper cognitive '
    'patterns\u2014including the aesthetic judgments that constitute scientific taste.'
)

doc.add_paragraph(
    'This review synthesizes these disparate threads to lay the groundwork for a new research '
    'program: computational modeling of individual scientists\' research taste. We organize our '
    'discussion around four key questions: (1) What is scientific taste, and how has it been '
    'conceptualized? (2) What computational tools exist for scientific hypothesis generation? '
    '(3) How can LLMs be personalized to capture individual cognitive profiles? (4) What are the '
    'key challenges and opportunities for formal taste modeling?'
)

# ── 2. The Concept of Taste ────────────────────────────────────────────────
doc.add_heading('2. The Concept of Taste in Science', level=1)

doc.add_heading('2.1 Philosophical Foundations', level=2)
doc.add_paragraph(
    'The philosophical treatment of taste begins with Immanuel Kant\'s Critique of Judgment (1790), '
    'which distinguished between determinate judgments (governed by rules) and reflective judgments '
    '(guided by aesthetic feeling). Kant argued that aesthetic judgment involves a "free play" of '
    'imagination and understanding that produces a sense of purposiveness without a definite purpose '
    '[13]. This Kantian framework proved remarkably prescient for understanding scientific taste: '
    'scientists often describe their theoretical preferences using aesthetic language\u2014"elegant," '
    '"beautiful," "natural"\u2014that resists reduction to explicit rules.'
)

doc.add_paragraph(
    'Pierre Bourdieu\'s sociological analysis of taste [14] added a crucial dimension: taste is not '
    'merely individual but is shaped by field-specific habitus\u2014the internalized dispositions '
    'acquired through participation in a scientific community. A physicist\'s taste is formed through '
    'years of exposure to exemplary work, peer evaluation, and methodological training. This social '
    'dimension has important implications for computational modeling: taste cannot be extracted from '
    'an individual\'s writings alone but must be understood in the context of their scientific community.'
)

doc.add_heading('2.2 Tacit Knowledge and Scientific Judgment', level=2)
doc.add_paragraph(
    'Michael Polanyi\'s Personal Knowledge (1958) [5] and The Tacit Dimension (1966) [15] provided '
    'the most influential philosophical framework for understanding the implicit nature of scientific '
    'expertise. Polanyi argued that "we know more than we can tell"\u2014that much of scientific '
    'judgment relies on subsidiary awareness that integrates disparate clues into a coherent '
    'perception of significance. For Polanyi, the selection of problems, the evaluation of evidence, '
    'and the recognition of theoretical beauty all involve tacit knowing that is embodied in the '
    'scientist\'s practice rather than explicitly formulated.'
)

doc.add_paragraph(
    'Polanyi\'s framework suggests that scientific taste operates at multiple levels: (a) perceptual '
    'taste\u2014the ability to notice patterns and anomalies in data; (b) methodological taste\u2014'
    'preferences for certain types of explanations, formalisms, or experimental approaches; '
    '(c) strategic taste\u2014judgments about which problems are important, tractable, and timely; '
    'and (d) aesthetic taste\u2014sensitivity to mathematical elegance, conceptual economy, and '
    'theoretical coherence. This multi-dimensional structure has direct implications for computational '
    'modeling through "taste axes."'
)

doc.add_heading('2.3 Thematic Analysis and Scientific Themata', level=2)
doc.add_paragraph(
    'Gerald Holton\'s thematic analysis [7] operationalized the concept of scientific taste by '
    'identifying specific "themata"\u2014persistent presuppositions that guide scientists\' choices '
    'across different domains and historical periods. Holton studied Einstein extensively, '
    'identifying core themata including symmetry, unity, simplicity, completeness, continuum, '
    'causality, and invariance [16]. These themata function as a scientist\'s implicit evaluation '
    'criteria, shaping which theories are considered promising, which anomalies are taken seriously, '
    'and which research directions are pursued.'
)

doc.add_paragraph(
    'Holton\'s approach is particularly valuable for computational modeling because it provides '
    'named, identifiable dimensions along which taste can be measured. Unlike Polanyi\'s wholly '
    'tacit knowledge, Holton\'s themata can be (partially) extracted from a scientist\'s published '
    'work, correspondence, and documented methodological choices. The challenge lies in determining '
    'the relative weights of different themata, how they interact, and how they evolve over time.'
)

doc.add_heading('2.4 The "AI Can Learn Scientific Taste" Perspective', level=2)
doc.add_paragraph(
    'A recent perspective piece titled "AI Can Learn Scientific Taste" (2024) [17] argued that '
    'modern AI systems, particularly large language models trained on scientific corpora, can '
    'internalize some aspects of scientific taste through exposure to the patterns of successful '
    'scientific inquiry. The authors distinguished between "shallow taste" (preference for certain '
    'stylistic or structural features) and "deep taste" (genuine understanding of what makes a '
    'scientific contribution significant). They argued that current LLMs primarily exhibit shallow '
    'taste but that deeper taste modeling is achievable through targeted training on expert '
    'evaluations and peer review data.'
)

# ── 3. AI for Scientific Discovery ─────────────────────────────────────────
doc.add_heading('3. AI Systems for Scientific Discovery', level=1)

doc.add_heading('3.1 The AI Scientist', level=2)
doc.add_paragraph(
    'The AI Scientist [8], developed by Sakana AI (Lu et al., 2024), represents the most ambitious '
    'attempt at fully automated scientific research. The system uses LLMs to perform the complete '
    'research cycle: generating research ideas, writing code to implement experiments, executing '
    'experiments, and writing scientific papers. Evaluated on machine learning subfields (diffusion '
    'modeling, language modeling, and grokking), The AI Scientist demonstrated the ability to produce '
    'papers that, while not matching expert quality, contained genuine novel contributions. '
    'Critically, the system lacks an explicit model of research taste\u2014it generates ideas broadly '
    'and relies on post-hoc filtering rather than taste-guided selection.'
)

doc.add_heading('3.2 MooseChem', level=2)
doc.add_paragraph(
    'MooseChem [18] (Li et al., 2024) addresses scientific hypothesis generation in chemistry, '
    'using a multi-agent framework where specialized LLM agents handle different aspects of the '
    'research process: literature survey, hypothesis formulation, and experimental design. The system '
    'demonstrates that domain-specific structure can improve LLM-based hypothesis generation. '
    'However, like The AI Scientist, MooseChem does not model the taste preferences of individual '
    'researchers\u2014it generates hypotheses without regard to whether they match a particular '
    'scientist\'s aesthetic or methodological preferences.'
)

doc.add_heading('3.3 SciMON and Scientific Inspiration', level=2)
doc.add_paragraph(
    'SciMON (Scientific Inspiration Machines Optimized for Novelty) [9] (Wang et al., 2023) takes '
    'a different approach, focusing on generating novel scientific ideas by optimizing for both '
    'relevance and novelty. The system retrieves relevant prior work and iteratively refines '
    'generated ideas to maximize their novelty relative to existing literature. SciMON\'s emphasis '
    'on novelty optimization is relevant to taste modeling because novelty is one dimension of '
    'scientific taste\u2014some scientists (like Einstein) privileged conceptual novelty, while others '
    'preferred incremental extension of established frameworks.'
)

doc.add_heading('3.4 Scideator', level=2)
doc.add_paragraph(
    'Scideator [19] (Radensky et al., 2024) is an interactive scientific ideation tool that helps '
    'researchers generate and refine research ideas through a structured process of facet '
    'recombination. Users specify "facets" of existing work (problems, methods, datasets, findings) '
    'and the system generates novel combinations. Scideator\'s facet-based approach bears resemblance '
    'to taste modeling: facet preferences could serve as a proxy for research taste, capturing which '
    'types of problems, methods, and approaches a scientist finds appealing.'
)

doc.add_heading('3.5 EvoScientist', level=2)
doc.add_paragraph(
    'EvoScientist [20] (Zhang et al., 2024) introduces evolutionary computation principles into '
    'LLM-based scientific discovery. The system maintains a population of research hypotheses that '
    'undergo mutation (LLM-based variation), crossover (idea combination), and selection (quality '
    'filtering). The evolutionary framework is relevant to taste modeling because selection pressure '
    'can be parameterized to reflect different taste profiles\u2014a "simplicity-biased" selection '
    'would favor Occam\'s-razor-style hypotheses, while a "unification-biased" selection would favor '
    'hypotheses that connect disparate phenomena.'
)

doc.add_heading('3.6 SciMuse', level=2)
doc.add_paragraph(
    'SciMuse [21] focuses on scientific inspiration by leveraging cross-domain knowledge transfer. '
    'The system identifies analogical connections between different scientific fields, suggesting '
    'that insights from one domain may be applicable to another. This cross-domain reasoning is '
    'directly relevant to taste modeling: Einstein\'s taste for unity, for instance, predisposed him '
    'to seek connections between seemingly separate physical phenomena (electromagnetism and gravity, '
    'energy and mass).'
)

doc.add_heading('3.7 Research Proposal Prediction', level=2)
doc.add_paragraph(
    '"Learning to Predict Future-Aligned Research Proposals with Language Models" [22] (Ye et al., '
    '2024) addresses the problem of predicting which research proposals will align with future '
    'developments in a field. The system trains on historical data to learn which types of proposals '
    'proved influential. This predictive approach complements taste modeling: while taste modeling '
    'asks "what would scientist X find promising?", proposal prediction asks "what will the field '
    'find important?" The intersection\u2014predicting individual scientists\' reactions to research '
    'proposals\u2014remains unexplored.'
)

# ── 4. LLM Personalization and Cognitive Simulation ─────────────────────────
doc.add_heading('4. LLM Personalization and Cognitive Simulation', level=1)

doc.add_heading('4.1 LaMP: Language Model Personalization', level=2)
doc.add_paragraph(
    'The LaMP benchmark [10] (Salemi et al., 2023) provides a systematic framework for evaluating '
    'personalized text generation. LaMP includes tasks such as personalized email subject generation, '
    'product review writing, and news headline generation, each requiring the model to adapt to '
    'individual users\' styles and preferences. The benchmark uses retrieval-augmented approaches '
    'where a user\'s historical outputs serve as context for personalizing new generations. LaMP\'s '
    'methodology is directly applicable to taste modeling: a scientist\'s published papers, reviews, '
    'and correspondence serve as the "user profile" from which taste preferences can be extracted.'
)

doc.add_heading('4.2 Cognitive Simulation Beyond Behavioral Imitation', level=2)
doc.add_paragraph(
    '"Can Large Language Models Simulate Human Cognition Beyond Behavioral Imitation?" [11] (Shu '
    'et al., 2024) critically examines whether LLMs can genuinely model human cognitive processes '
    'or merely reproduce surface-level behavioral patterns. The authors argue that behavioral '
    'imitation (matching human outputs) is insufficient for true cognitive simulation, which '
    'requires modeling the underlying processes\u2014attention, memory, reasoning strategies\u2014that '
    'produce those outputs. For taste modeling, this distinction is crucial: reproducing Einstein\'s '
    'theory preferences is easier than modeling the aesthetic judgment process that generated them.'
)

doc.add_paragraph(
    'The paper identifies several key challenges: (a) the underspecification problem\u2014many '
    'different cognitive processes can produce the same behavioral output; (b) the training data '
    'gap\u2014LLMs are trained on text outputs, not cognitive processes; (c) the evaluation problem\u2014'
    'how to verify that a model captures genuine cognitive patterns rather than surface correlations. '
    'These challenges directly apply to scientific taste modeling and motivate our approach of '
    'grounding taste models in historical evidence rather than pure behavioral fitting.'
)

doc.add_heading('4.3 MirrorMind: Personality and Cognitive Modeling', level=2)
doc.add_paragraph(
    'MirrorMind [12] (2024) attempts to create LLM-based "cognitive mirrors" that reflect '
    'individual users\' thinking patterns, personality traits, and decision-making styles. The '
    'system uses structured self-report data and behavioral observation to construct cognitive '
    'profiles, then fine-tunes LLMs to generate responses consistent with those profiles. While '
    'MirrorMind focuses on general cognitive traits rather than scientific taste specifically, its '
    'methodology\u2014combining explicit trait measurement with behavioral data\u2014provides a useful '
    'template for taste modeling.'
)

# ── 5. Gaps and Synthesis ──────────────────────────────────────────────────
doc.add_heading('5. Critical Gaps and Synthesis', level=1)

doc.add_paragraph(
    'Our review reveals several critical gaps at the intersection of these research streams:'
)

p = doc.add_paragraph()
p.style = doc.styles['Normal']
run = p.add_run('Gap 1: No Formal Model of Scientific Taste. ')
run.bold = True
p.add_run(
    'Despite extensive philosophical discussion, no computational model exists that formally '
    'defines scientific taste as a measurable, multi-dimensional construct. Holton\'s themata '
    'provide the closest approximation, but they lack quantitative specification of weights, '
    'interactions, and temporal dynamics.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 2: Individual vs. Community Taste. ')
run.bold = True
p.add_run(
    'Existing AI-for-science systems (AI Scientist, MooseChem, SciMON) optimize for generic '
    'scientific quality rather than individual researchers\' preferences. No system models how '
    'specific scientists\' taste profiles differ from community norms.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 3: Evidence Grounding. ')
run.bold = True
p.add_run(
    'LLM-based cognitive simulation (MirrorMind) and personalization (LaMP) rely on behavioral '
    'fitting without requiring historical evidence to support their models. For historical '
    'scientists, where ground truth is unavailable, evidence grounding becomes essential to '
    'distinguish model hallucination from genuine taste modeling.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 4: Temporal Dynamics. ')
run.bold = True
p.add_run(
    'Scientists\' taste evolves over their careers (e.g., Einstein\'s shift from physical '
    'intuition to mathematical strategy around 1920 [3]). No existing system models this '
    'temporal evolution or enforces temporal cutoffs to prevent anachronism.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 5: Evaluation Methodology. ')
run.bold = True
p.add_run(
    'No benchmark exists for evaluating whether a computational model accurately captures a '
    'specific scientist\'s taste. Existing benchmarks evaluate hypothesis quality (SciMON) or '
    'personalization accuracy (LaMP) but not taste fidelity.'
)

# ── 6. Toward Computational Taste Modeling ──────────────────────────────────
doc.add_heading('6. Toward Computational Research Taste Modeling', level=1)

doc.add_paragraph(
    'Based on our review, we propose a framework for computational research taste modeling that '
    'addresses the identified gaps. The key elements are:'
)

doc.add_paragraph(
    '(1) Taste Axes: Following Holton\'s thematic analysis, we define taste as a vector in a '
    'multi-dimensional space of scientific preferences (e.g., simplicity, unity, invariance, '
    'empirical grounding, mathematical beauty). Each axis is grounded in documented evidence.'
)
doc.add_paragraph(
    '(2) Evidence-Based Scoring: Taste axis weights are derived from historical evidence '
    '(primary sources, scholarly analyses) rather than behavioral fitting, ensuring transparency '
    'and preventing hallucination.'
)
doc.add_paragraph(
    '(3) Temporal Dynamics: The model supports temporal cutoffs and period-specific weight '
    'adjustments to capture how a scientist\'s taste evolves over their career.'
)
doc.add_paragraph(
    '(4) LLM-RAG Pipeline: Retrieval-augmented generation combines evidence retrieval with '
    'LLM-based evaluation, grounding each score in specific historical evidence.'
)
doc.add_paragraph(
    '(5) Explicit Inference Marking: The system strictly distinguishes between what is supported '
    'by evidence and what is model inference, addressing the transparency gap.'
)

# ── 7. Conclusion ──────────────────────────────────────────────────────────
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    'Scientific research taste represents a critical yet computationally underexplored dimension '
    'of scientific cognition. While philosophers have long recognized its importance (Polanyi, '
    'Holton, Kuhn) and AI researchers have built increasingly capable scientific discovery systems '
    '(AI Scientist, SciMON, MooseChem), the specific challenge of modeling individual scientists\' '
    'taste profiles remains open. The convergence of LLM capabilities, personalization methods, '
    'and rich historical scholarship on eminent scientists creates a unique opportunity to bridge '
    'this gap. We believe that evidence-based, multi-dimensional taste modeling\u2014grounded in '
    'historical sources and transparent about its limitations\u2014represents a promising direction '
    'for both AI-assisted scientific discovery and history/philosophy of science.'
)

# ── References ──────────────────────────────────────────────────────────────
doc.add_heading('References', level=1)

refs = [
    '[1] Feynman, R. P. (1974). "Cargo Cult Science." Caltech Commencement Address.',
    '[2] Pais, A. (1982). Subtle is the Lord: The Science and the Life of Albert Einstein. Oxford University Press.',
    '[3] van Dongen, J. (2010). Einstein\'s Unification. Cambridge University Press.',
    '[4] Medawar, P. B. (1967). The Art of the Soluble. Methuen.',
    '[5] Polanyi, M. (1958). Personal Knowledge: Towards a Post-Critical Philosophy. University of Chicago Press.',
    '[6] Kuhn, T. S. (1962). The Structure of Scientific Revolutions. University of Chicago Press.',
    '[7] Holton, G. (1973). Thematic Origins of Scientific Thought: Kepler to Einstein. Harvard University Press.',
    '[8] Lu, C., et al. (2024). "The AI Scientist: Towards Fully Automated Open-Ended Scientific Discovery." arXiv:2408.06292.',
    '[9] Wang, Q., et al. (2023). "SciMON: Scientific Inspiration Machines Optimized for Novelty." arXiv:2305.14259.',
    '[10] Salemi, A., et al. (2023). "LaMP: When Large Language Models Meet Personalization." arXiv:2304.11406.',
    '[11] Shu, T., et al. (2024). "Can Large Language Models Simulate Human Cognition Beyond Behavioral Imitation?" arXiv.',
    '[12] MirrorMind. (2024). LLM-based Cognitive Mirror for Personality Modeling. arXiv.',
    '[13] Kant, I. (1790). Critique of Judgment. Trans. Werner Pluhar, Hackett Publishing.',
    '[14] Bourdieu, P. (1984). Distinction: A Social Critique of the Judgement of Taste. Harvard University Press.',
    '[15] Polanyi, M. (1966). The Tacit Dimension. University of Chicago Press.',
    '[16] Holton, G. (1998). The Scientific Imagination. Harvard University Press.',
    '[17] "AI Can Learn Scientific Taste." (2024). Perspective article.',
    '[18] Li, Z., et al. (2024). "MooseChem: Large Language Models for Rediscovering Unseen Chemistry Scientific Hypotheses." arXiv.',
    '[19] Radensky, M., et al. (2024). "Scideator: Human-LLM Scientific Idea Generation Grounded in Research-Paper Facet Recombination." arXiv.',
    '[20] Zhang, C., et al. (2024). "EvoScientist: Evolving Scientific Hypotheses with Language Models." arXiv.',
    '[21] SciMuse. (2024). Cross-Domain Scientific Inspiration System.',
    '[22] Ye, J., et al. (2024). "Learning to Predict Future-Aligned Research Proposals with Language Models." arXiv.',
    '[23] Einstein, A. (1933). "On the Method of Theoretical Physics." Herbert Spencer Lecture, Oxford.',
    '[24] Einstein, A. (1949). "Autobiographical Notes." In P.A. Schilpp (ed.), Albert Einstein: Philosopher-Scientist.',
    '[25] Fine, A. (1986). The Shaky Game: Einstein, Realism, and the Quantum Theory. University of Chicago Press.',
    '[26] Norton, J. (1984). "How Einstein Found His Field Equations." Historical Studies in the Physical Sciences.',
    '[27] Howard, D. (1985). "Einstein on Locality and Separability." Studies in History and Philosophy of Science.',
    '[28] Isaacson, W. (2007). Einstein: His Life and Universe. Simon & Schuster.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ────────────────────────────────────────────────────────────────────
output_dir = r"D:\student\scientist_taste\EinsteinResearchTaste\docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "literature_review.docx")
doc.save(output_path)
print(f"Saved literature review to {output_path}")
