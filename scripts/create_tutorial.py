"""Generate the Skill Publishing Tutorial Word document (bilingual)."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.4
style.paragraph_format.space_after = Pt(4)

for level, size in [('Heading 1', 16), ('Heading 2', 13), ('Heading 3', 11)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 51, 102)

# ── Title ────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Research Taste Skill Publishing Tutorial\n')
run.font.size = Pt(20)
run.bold = True
run = title.add_run('How to Package and Publish Einstein/Feynman Taste as Claude Code Skills\n')
run.font.size = Pt(13)
run.italic = True
run = title.add_run('\n')
run = title.add_run('Version 1.0 | April 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Overview / 概述',
    '2. Prerequisites / 前置要求',
    '3. Understanding Claude Code Skills / 理解Claude Code Skills',
    '4. Step-by-Step: Packaging as a Skill / 逐步打包为Skill',
    '5. Step-by-Step: Publishing to Community / 发布到社区',
    '6. Testing Your Published Skill / 测试已发布的Skill',
    '7. Maintenance and Updates / 维护和更新',
    '8. FAQ / 常见问题',
]
for item in toc:
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('1. Overview / 概述', level=1)

doc.add_paragraph(
    'This tutorial explains how to package the Einstein Research Taste and Feynman Research '
    'Taste modeling systems as Claude Code Skills, and how to publish them to the Claude Code '
    'community for others to use.'
)
doc.add_paragraph(
    '本教程说明如何将 Einstein Research Taste 和 Feynman Research Taste 建模系统打包为 '
    'Claude Code Skills，并发布到 Claude Code 社区供他人使用。'
)

doc.add_heading('What is a Claude Code Skill?', level=2)
doc.add_paragraph(
    'A Claude Code Skill is a reusable capability that extends Claude Code (the AI coding '
    'assistant). Skills are defined by a SKILL.md file that tells Claude how to use the '
    'capability. When a user invokes a skill, Claude reads the SKILL.md instructions and '
    'executes them in the user\'s environment.'
)
doc.add_paragraph(
    'Claude Code Skill 是一种可复用的能力扩展，用于增强 Claude Code（AI编码助手）。'
    'Skills 由一个 SKILL.md 文件定义，该文件告诉 Claude 如何使用该能力。'
    '当用户调用 Skill 时，Claude 读取 SKILL.md 中的指令并在用户环境中执行。'
)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('2. Prerequisites / 前置要求', level=1)

doc.add_heading('2.1 Software Requirements / 软件要求', level=2)
p = doc.add_paragraph()
items = [
    'Claude Code CLI installed (npm install -g @anthropic-ai/claude-code)',
    'Python 3.10+ installed',
    'Git installed',
    'GitHub account (for community publishing)',
    'Anthropic API key (for LLM-based evaluation)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Project Requirements / 项目要求', level=2)
doc.add_paragraph(
    'Ensure both projects are installed and tests pass:'
)
p = doc.add_paragraph()
run = p.add_run(
    'cd EinsteinResearchTaste && pip install -e . && python -m pytest tests/ -q\n'
    'cd FeynmanResearchTaste && pip install -e . && python -m pytest tests/ -q'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('3. Understanding Claude Code Skills / 理解Claude Code Skills', level=1)

doc.add_heading('3.1 Skill File Structure / Skill文件结构', level=2)
doc.add_paragraph(
    'A Claude Code skill consists of a single SKILL.md file placed in a specific directory. '
    'The file has two parts:'
)
doc.add_paragraph('Part 1: YAML Frontmatter (metadata)', style='List Bullet')
doc.add_paragraph('Part 2: Markdown body (instructions for Claude)', style='List Bullet')

doc.add_paragraph('Example structure / 示例结构:')
p = doc.add_paragraph()
run = p.add_run(
    'your-project/\n'
    '  .claude/\n'
    '    skills/\n'
    '      your-skill-name/\n'
    '        SKILL.md          <-- The skill definition file\n'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('3.2 SKILL.md Format / SKILL.md 格式', level=2)
doc.add_paragraph('The SKILL.md file format:')
p = doc.add_paragraph()
run = p.add_run(
    '---\n'
    'name: your-skill-name\n'
    'description: One-line description of what the skill does\n'
    '---\n'
    '\n'
    '# Skill Title\n'
    '\n'
    'Detailed instructions for Claude on how to use this skill.\n'
    'Include: what it does, when to use it, code examples, API reference.\n'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('3.3 How Skills Are Discovered / Skill如何被发现', level=2)
doc.add_paragraph(
    'Claude Code discovers skills in two ways:\n'
    '1. Local skills: From .claude/skills/ in the current project directory\n'
    '2. Community skills: Published via Claude Code\'s skill registry\n'
    '\n'
    'Claude Code 通过两种方式发现 Skills：\n'
    '1. 本地技能：从当前项目目录的 .claude/skills/ 中\n'
    '2. 社区技能：通过 Claude Code 的技能注册表发布'
)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('4. Step-by-Step: Packaging as a Skill / 逐步打包', level=1)

doc.add_heading('Step 1: Verify Your SKILL.md / 验证SKILL.md文件', level=2)
doc.add_paragraph('Both projects already have SKILL.md files:')
p = doc.add_paragraph()
run = p.add_run(
    'EinsteinResearchTaste/.claude/skills/einstein-taste/SKILL.md\n'
    'FeynmanResearchTaste/.claude/skills/feynman-taste/SKILL.md'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph(
    'Check that your SKILL.md contains:\n'
    '  - Valid YAML frontmatter with name and description\n'
    '  - Clear instructions on what the skill does\n'
    '  - Code examples showing how to invoke the skill\n'
    '  - List of available functions/capabilities\n'
    '\n'
    '确认 SKILL.md 包含：\n'
    '  - 有效的 YAML frontmatter（name 和 description）\n'
    '  - 清晰的技能用途说明\n'
    '  - 代码示例展示如何调用\n'
    '  - 可用功能列表'
)

doc.add_heading('Step 2: Create a GitHub Repository / 创建GitHub仓库', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '# For Einstein project\n'
    'cd EinsteinResearchTaste\n'
    'git init\n'
    'git add .\n'
    'git commit -m "Initial commit: Einstein Research Taste v0.1.0"\n'
    'gh repo create einstein-research-taste --public --source=. --push\n'
    '\n'
    '# For Feynman project\n'
    'cd FeynmanResearchTaste\n'
    'git init\n'
    'git add .\n'
    'git commit -m "Initial commit: Feynman Research Taste v0.1.0"\n'
    'gh repo create feynman-research-taste --public --source=. --push'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('Step 3: Ensure Dependencies Are Documented / 确保依赖已记录', level=2)
doc.add_paragraph(
    'The pyproject.toml already lists all dependencies. Users will need to install '
    'the package before using the skill. Add installation instructions to your SKILL.md.'
)
doc.add_paragraph(
    'pyproject.toml 已列出所有依赖。用户在使用 Skill 前需要安装包。'
    '请在 SKILL.md 中添加安装说明。'
)

doc.add_heading('Step 4: Test Locally / 本地测试', level=2)
doc.add_paragraph('Test that the skill works in Claude Code:')
p = doc.add_paragraph()
run = p.add_run(
    '# Open Claude Code in the project directory\n'
    'cd EinsteinResearchTaste\n'
    'claude\n'
    '\n'
    '# In Claude Code, try:\n'
    '> Use the einstein-taste skill to evaluate "unified field theory" at year 1935\n'
    '\n'
    '# Claude should read the SKILL.md and execute the skill'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('5. Step-by-Step: Publishing to Community / 发布到社区', level=1)

doc.add_heading('5.1 Method A: Via Claude Code CLI / 通过CLI发布', level=2)
doc.add_paragraph(
    'Claude Code provides a built-in mechanism for sharing skills. As of early 2026, '
    'the primary method is through the Claude Code skill registry.'
)
p = doc.add_paragraph()
run = p.add_run(
    '# Step 1: Navigate to your project\n'
    'cd EinsteinResearchTaste\n'
    '\n'
    '# Step 2: Use the /skill command in Claude Code to manage skills\n'
    'claude\n'
    '> /skill publish einstein-taste\n'
    '\n'
    '# Follow the prompts to:\n'
    '#   - Confirm the skill name and description\n'
    '#   - Set visibility (public/private)\n'
    '#   - Add tags for discoverability\n'
    '#   - Provide a GitHub repository URL'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('5.2 Method B: Via GitHub + Manual Registration / GitHub + 手动注册', level=2)
doc.add_paragraph(
    'If the CLI publishing is not available, you can share your skill by:'
)
items_pub = [
    'Push your project to a public GitHub repository',
    'Create a clear README with installation and usage instructions',
    'Add the .claude/skills/ directory with your SKILL.md',
    'Tag a release (git tag v0.1.0 && git push --tags)',
    'Share the repository URL in Claude Code community channels',
    'Users can then clone your repo and the skill will be auto-discovered',
]
for item in items_pub:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Method C: Using the Skill Creator / 使用Skill Creator', level=2)
doc.add_paragraph(
    'Claude Code has a built-in skill creator (anthropic-skills:skill-creator) that can '
    'help you create, test, and optimize skills. Use it like this:'
)
p = doc.add_paragraph()
run = p.add_run(
    '# In Claude Code:\n'
    '> /skill-creator create a skill that evaluates scientific theories \\\n'
    '  against Einstein\'s research taste using the einstein_taste Python package\n'
    '\n'
    '# The skill creator will:\n'
    '#   - Generate an optimized SKILL.md\n'
    '#   - Run eval tests to measure skill quality\n'
    '#   - Suggest improvements'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('6. Testing Your Published Skill / 测试已发布的Skill', level=1)

doc.add_heading('6.1 Install from Repository / 从仓库安装', level=2)
p = doc.add_paragraph()
run = p.add_run(
    '# Clone and install\n'
    'git clone https://github.com/YOUR_USERNAME/einstein-research-taste.git\n'
    'cd einstein-research-taste\n'
    'pip install -e .\n'
    '\n'
    '# Open Claude Code and test\n'
    'claude\n'
    '> Evaluate "quantum field theory" using Einstein\'s research taste at year 1935'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('6.2 Verification Checklist / 验证清单', level=2)
checks = [
    'Skill is discovered by Claude Code (appears in /skills list)',
    'Skill executes correctly when invoked',
    'Output includes evidence-based scores with [EVIDENCE]/[INFERRED] labels',
    'Temporal cutoff works (different results for different years)',
    'Ranking produces correct ordering for benchmark cases',
    'Error messages are clear when API key is missing',
    'Installation instructions in SKILL.md are accurate and complete',
]
for check in checks:
    doc.add_paragraph(check, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('7. Maintenance and Updates / 维护和更新', level=1)

doc.add_heading('7.1 Updating Evidence / 更新证据', level=2)
doc.add_paragraph(
    'To add new evidence records, edit the seed_evidence.py file or run the enrichment '
    'script. Then bump the version in pyproject.toml and push.'
)

doc.add_heading('7.2 Adding New Scientists / 添加新科学家', level=2)
doc.add_paragraph(
    'To create a taste model for a new scientist (e.g., Newton, Mendel):\n'
    '1. Copy the project structure from Einstein or Feynman\n'
    '2. Define taste axes specific to that scientist in config/settings.py\n'
    '3. Define temporal periods matching their career\n'
    '4. Create seed evidence from their writings and scholarly analyses\n'
    '5. Create benchmark cases from documented preferences\n'
    '6. Update SKILL.md with the new scientist\'s information\n'
    '7. Run tests and verify the pipeline'
)

doc.add_heading('7.3 Versioning / 版本管理', level=2)
doc.add_paragraph(
    'Follow semantic versioning: MAJOR.MINOR.PATCH\n'
    '- PATCH: Bug fixes, evidence corrections\n'
    '- MINOR: New evidence, new benchmark cases, improved retrieval\n'
    '- MAJOR: New taste axes, architecture changes, breaking API changes'
)

# ══════════════════════════════════════════════════════════════════════
doc.add_heading('8. FAQ / 常见问题', level=1)

faqs = [
    ('Q: Do users need an API key to use the skill?\n'
     'A: No. The offline mode works without an API key using keyword-based scoring. '
     'For full LLM-based evaluation, users need an Anthropic or OpenAI API key.\n\n'
     'Q: 用户是否需要API密钥？\n'
     'A: 不需要。离线模式使用关键词评分。完整的LLM评估需要API密钥。'),

    ('Q: Can I use a different LLM provider?\n'
     'A: Yes. The evaluator supports both Anthropic (Claude) and OpenAI (GPT-4). '
     'Change the llm_provider setting in ModelConfig.\n\n'
     'Q: 可以使用其他LLM吗？\n'
     'A: 可以。评估器支持 Anthropic (Claude) 和 OpenAI (GPT-4)。修改 ModelConfig 中的 llm_provider。'),

    ('Q: How accurate is the taste model?\n'
     'A: The model is grounded in historical evidence and designed for scholarly use. '
     'Benchmark accuracy depends on evidence quality and LLM capability. Always check '
     'the [EVIDENCE]/[INFERRED] labels in the output.\n\n'
     'Q: 品味模型有多准确？\n'
     'A: 模型基于历史证据，为学术用途设计。始终检查输出中的[EVIDENCE]/[INFERRED]标签。'),

    ('Q: Can this be used for role-playing?\n'
     'A: No. This system explicitly avoids role-playing. It evaluates theories against '
     'documented taste axes, not by pretending to be a historical figure.\n\n'
     'Q: 可以用于角色扮演吗？\n'
     'A: 不能。系统明确避免角色扮演。它基于已记录的品味轴评估理论，而非假装是历史人物。'),
]

for faq in faqs:
    doc.add_paragraph(faq)

# ── Save ────────────────────────────────────────────────────────────
output_dir = r"D:\student\scientist_taste\EinsteinResearchTaste\docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "skill_publishing_tutorial.docx")
doc.save(output_path)

# Also save a copy to Feynman project
output_path2 = r"D:\student\scientist_taste\FeynmanResearchTaste\docs"
os.makedirs(output_path2, exist_ok=True)
import shutil
shutil.copy(output_path, os.path.join(output_path2, "skill_publishing_tutorial.docx"))

print(f"Tutorial saved to {output_path}")
print(f"Copy saved to {output_path2}")
