"""
Update the literature review with additional references from research agents.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 12, True)]:
    s = doc.styles[level]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Modeling Scientific Research Taste:\nA Comprehensive Literature Review')
run.font.size = Pt(18)
run.bold = True
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('From Cognitive Aesthetics to LLM-Based Scientific Discovery')
run.font.size = Pt(14)
run.italic = True
run.font.name = 'Times New Roman'

doc.add_paragraph()

# ── Abstract ────────────────────────────────────────────────────────────────
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Scientific research taste\u2014the implicit preferences, aesthetic judgments, and methodological '
    'inclinations that guide scientists in selecting problems, evaluating theories, and directing '
    'inquiry\u2014has long been recognized as a crucial yet elusive component of scientific excellence. '
    'This review traces the concept of taste from its philosophical roots in Kantian aesthetics and '
    'Polanyi\'s tacit knowledge through modern cognitive science frameworks, culminating in recent '
    'computational approaches using large language models (LLMs). We survey four interconnected '
    'research streams: (1) the philosophical and cognitive science foundations of scientific taste; '
    '(2) AI systems for autonomous scientific discovery and hypothesis generation; '
    '(3) LLM-based personalization and cognitive simulation; and '
    '(4) computational aesthetics and preference modeling. We identify key gaps in the literature, '
    'particularly the absence of formal computational models that capture individual scientists\' '
    'research taste profiles, and propose directions for evidence-based taste modeling that bridges '
    'historical scholarship with modern AI capabilities.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'The notion that scientific excellence involves a form of "taste" has deep roots in the history '
    'and philosophy of science. Richard Feynman famously described the ability to identify '
    'promising research directions as requiring a kind of scientific taste that cannot be taught '
    'through textbooks alone [1]. Albert Einstein\'s persistent pursuit of unified field theory, '
    'despite decades of failure, reflected aesthetic commitments\u2014to simplicity, symmetry, and '
    'mathematical beauty\u2014that transcended purely empirical considerations [2, 3]. Peter Medawar '
    'argued that the selection of a scientific problem is itself a creative act guided by an '
    '"aesthetic sense" for what is important and tractable [4].'
)

doc.add_paragraph(
    'Despite its acknowledged importance, scientific taste has resisted formal characterization. '
    'Michael Polanyi\'s concept of "tacit knowledge" [5] captured the intuition that much of what '
    'guides expert scientific judgment is difficult to articulate explicitly. Thomas Kuhn\'s notion '
    'of "paradigm-dependent values" [6] highlighted how aesthetic criteria shift across scientific '
    'revolutions. Gerald Holton\'s "thematic analysis" [7] identified persistent presuppositions '
    '(themata) that constrain scientists\' theoretical choices\u2014symmetry, unity, simplicity\u2014which '
    'function as taste axes guiding inquiry.'
)

doc.add_paragraph(
    'The recent revolution in large language models (LLMs) has opened new possibilities for '
    'computational approaches to scientific taste. Systems like The AI Scientist [8] and '
    'SciMON [9] demonstrate that LLMs can generate novel scientific hypotheses, while '
    'personalization benchmarks like LaMP [10] show that LLMs can be adapted to capture individual '
    'preferences. A landmark study by Si et al. [11] demonstrated that LLM-generated research ideas '
    'were rated as more novel (though slightly less feasible) than those of 100+ human NLP researchers, '
    'raising fundamental questions about the nature of scientific judgment. Concurrently, work on '
    'LLM-based cognitive simulation [12, 13] and scientist persona modeling [14, 15] raises the '
    'question of whether these models can go beyond behavioral mimicry to capture deeper cognitive '
    'patterns\u2014including the aesthetic judgments that constitute scientific taste.'
)

doc.add_paragraph(
    'This review synthesizes these disparate threads to lay the groundwork for a new research '
    'program: computational modeling of individual scientists\' research taste. We organize our '
    'discussion around five key questions: (1) What is scientific taste, and how has it been '
    'conceptualized? (2) What computational tools exist for scientific hypothesis generation? '
    '(3) How can LLMs be personalized to capture individual cognitive profiles? '
    '(4) What frameworks from computational aesthetics and preference learning apply? '
    '(5) What are the key challenges and opportunities for formal taste modeling?'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. THE CONCEPT OF TASTE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. The Concept of Taste in Science', level=1)

doc.add_heading('2.1 Philosophical Foundations', level=2)
doc.add_paragraph(
    'The philosophical treatment of taste begins with Immanuel Kant\'s Critique of Judgment (1790), '
    'which distinguished between determinate judgments (governed by rules) and reflective judgments '
    '(guided by aesthetic feeling). Kant argued that aesthetic judgment involves a "free play" of '
    'imagination and understanding that produces a sense of purposiveness without a definite purpose '
    '[16]. This Kantian framework proved remarkably prescient for understanding scientific taste: '
    'scientists often describe their theoretical preferences using aesthetic language\u2014"elegant," '
    '"beautiful," "natural"\u2014that resists reduction to explicit rules.'
)

doc.add_paragraph(
    'Pierre Bourdieu\'s sociological analysis of taste [17] added a crucial dimension: taste is not '
    'merely individual but is shaped by field-specific habitus\u2014the internalized dispositions '
    'acquired through participation in a scientific community. A physicist\'s taste is formed through '
    'years of exposure to exemplary work, peer evaluation, and methodological training. This social '
    'dimension has important implications for computational modeling: taste cannot be extracted from '
    'an individual\'s writings alone but must be understood in the context of their scientific community.'
)

doc.add_heading('2.2 Tacit Knowledge and Scientific Judgment', level=2)
doc.add_paragraph(
    'Michael Polanyi\'s Personal Knowledge (1958) [5] and The Tacit Dimension (1966) [18] provided '
    'the most influential philosophical framework for understanding the implicit nature of scientific '
    'expertise. Polanyi argued that "we know more than we can tell"\u2014that much of scientific '
    'judgment relies on subsidiary awareness that integrates disparate clues into a coherent '
    'perception of significance. For Polanyi, the selection of problems, the evaluation of evidence, '
    'and the recognition of theoretical beauty all involve tacit knowing that is embodied in the '
    'scientist\'s practice rather than explicitly formulated.'
)

doc.add_paragraph(
    'Polanyi\'s framework suggests that scientific taste operates at multiple levels: (a) perceptual '
    'taste\u2014the ability to notice patterns and anomalies in data; (b) methodological taste\u2014'
    'preferences for certain types of explanations, formalisms, or experimental approaches; '
    '(c) strategic taste\u2014judgments about which problems are important, tractable, and timely; '
    'and (d) aesthetic taste\u2014sensitivity to mathematical elegance, conceptual economy, and '
    'theoretical coherence. This multi-dimensional structure directly motivates the "taste axes" '
    'approach to computational modeling.'
)

doc.add_heading('2.3 Thematic Analysis and Scientific Themata', level=2)
doc.add_paragraph(
    'Gerald Holton\'s thematic analysis [7] operationalized the concept of scientific taste by '
    'identifying specific "themata"\u2014persistent presuppositions that guide scientists\' choices '
    'across different domains and historical periods. Holton studied Einstein extensively, '
    'identifying core themata including symmetry, unity, simplicity, completeness, continuum, '
    'causality, and invariance [19]. These themata function as a scientist\'s implicit evaluation '
    'criteria, shaping which theories are considered promising, which anomalies are taken seriously, '
    'and which research directions are pursued.'
)

doc.add_paragraph(
    'Holton\'s approach is particularly valuable for computational modeling because it provides '
    'named, identifiable dimensions along which taste can be measured. Unlike Polanyi\'s wholly '
    'tacit knowledge, Holton\'s themata can be (partially) extracted from a scientist\'s published '
    'work, correspondence, and documented methodological choices. The challenge lies in determining '
    'the relative weights of different themata, how they interact, and how they evolve over time.'
)

doc.add_heading('2.4 The "AI Can Learn Scientific Taste" Perspective', level=2)
doc.add_paragraph(
    'A recent perspective piece titled "AI Can Learn Scientific Taste" (2024) [20] argued that '
    'modern AI systems, particularly large language models trained on scientific corpora, can '
    'internalize some aspects of scientific taste through exposure to the patterns of successful '
    'scientific inquiry. The authors distinguished between "shallow taste" (preference for certain '
    'stylistic or structural features) and "deep taste" (genuine understanding of what makes a '
    'scientific contribution significant). They argued that current LLMs primarily exhibit shallow '
    'taste but that deeper taste modeling is achievable through targeted training on expert '
    'evaluations and peer review data.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. AI FOR SCIENTIFIC DISCOVERY
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. AI Systems for Scientific Discovery and Hypothesis Generation', level=1)

doc.add_heading('3.1 The AI Scientist', level=2)
doc.add_paragraph(
    'The AI Scientist [8], developed by Sakana AI (Lu et al., 2024), represents the most ambitious '
    'attempt at fully automated scientific research. The system uses LLMs (Claude/GPT-4) to perform '
    'the complete research cycle: generating research ideas, writing code to implement experiments, '
    'executing experiments, writing scientific papers, and even conducting peer review. Evaluated on '
    'machine learning subfields (diffusion modeling, language modeling, and grokking), The AI Scientist '
    'demonstrated the ability to produce papers that, while not matching expert quality, contained '
    'genuine novel contributions. Critically, the system lacks an explicit model of research taste\u2014'
    'it generates ideas broadly and relies on post-hoc filtering rather than taste-guided selection.'
)

doc.add_heading('3.2 MooseChem', level=2)
doc.add_paragraph(
    'MooseChem [21] (Li et al., 2024) addresses scientific hypothesis generation in chemistry, '
    'using a multi-agent framework where specialized LLM agents handle different aspects of the '
    'research process. The system simulates chemist reasoning by having LLMs rediscover known but '
    'withheld chemical hypotheses from the literature. This "rediscovery" evaluation paradigm is '
    'particularly relevant to taste modeling: one could test whether a taste model correctly '
    'identifies the hypotheses that a specific scientist found compelling.'
)

doc.add_heading('3.3 SciMON and Scientific Inspiration', level=2)
doc.add_paragraph(
    'SciMON (Scientific Inspiration Machines Optimized for Novelty) [9] (Radensky et al., 2023, '
    'ACL 2023) focuses on generating novel scientific ideas by optimizing for both relevance and '
    'novelty. The system retrieves relevant prior work and iteratively refines generated ideas to '
    'maximize their novelty relative to existing literature. SciMON\'s emphasis on novelty '
    'optimization is relevant to taste modeling because novelty is one dimension of scientific '
    'taste\u2014some scientists privileged conceptual novelty, while others preferred incremental extension.'
)

doc.add_heading('3.4 Scideator', level=2)
doc.add_paragraph(
    'Scideator [22] (Radensky et al., 2024) is an interactive scientific ideation tool that helps '
    'researchers generate and refine research ideas through structured facet recombination. Users '
    'specify "facets" of existing work (problems, methods, datasets, findings) and the system '
    'generates novel combinations. Scideator\'s facet-based approach bears resemblance to taste '
    'modeling: facet preferences could serve as a proxy for research taste.'
)

doc.add_heading('3.5 EvoScientist', level=2)
doc.add_paragraph(
    'EvoScientist [23] (Zhang et al., 2024) introduces evolutionary computation principles into '
    'LLM-based scientific discovery. The system maintains a population of research hypotheses that '
    'undergo mutation (LLM-based variation), crossover (idea combination), and selection (quality '
    'filtering). The evolutionary framework is relevant to taste modeling because selection pressure '
    'can be parameterized to reflect different taste profiles.'
)

doc.add_heading('3.6 SciMuse and Cross-Domain Inspiration', level=2)
doc.add_paragraph(
    'SciMuse [24] focuses on scientific inspiration by leveraging cross-domain knowledge transfer. '
    'The system identifies analogical connections between different scientific fields, suggesting '
    'that insights from one domain may apply to another. This cross-domain reasoning is directly '
    'relevant to taste: Einstein\'s taste for unity predisposed him to seek connections between '
    'seemingly separate phenomena (electromagnetism and gravity, energy and mass).'
)

doc.add_heading('3.7 Additional Systems', level=2)
doc.add_paragraph(
    'Several other systems merit mention. ResearchAgent [25] (Baek et al., 2024) implements an '
    'iterative pipeline that reviews literature, identifies gaps, generates hypotheses, and refines '
    'via feedback using the Semantic Scholar API. SciAgents [26] (Ghafarollahi & Buehler, 2024) '
    'uses multi-agent systems where different LLM agents play roles (ontologist, scientist, critic) '
    'for hypothesis generation in materials science using knowledge graphs. ChemReasoner [27] '
    '(Sprueill et al., 2024, ICML) combines Monte Carlo Tree Search with LLM reasoning for '
    'catalyst discovery. VirtualLab [28] (Swanson et al., 2024) creates a multi-persona virtual '
    'lab with agents playing PI, immunologist, and computational biologist roles.'
)

doc.add_heading('3.8 Idea Quality Evaluation', level=2)
doc.add_paragraph(
    'A landmark study by Si et al. [11] (2024) had GPT-4 and Claude generate research ideas, '
    'which were then blind-evaluated by 100+ NLP researchers alongside human expert ideas. LLM '
    'ideas were rated as more novel but slightly less feasible. This established a key benchmark '
    'methodology for evaluating AI-generated scientific ideas and implicitly raised the question '
    'of whether "novelty preference" constitutes a form of scientific taste that LLMs exhibit. '
    'Liang et al. [29] (2024) further demonstrated significant overlap between LLM and human '
    'peer reviews, suggesting LLMs capture some aspects of scientific judgment.'
)

doc.add_heading('3.9 Future-Aligned Proposal Prediction', level=2)
doc.add_paragraph(
    '"Learning to Predict Future-Aligned Research Proposals with Language Models" [30] (Ye et al., '
    '2024) addresses predicting which research proposals will align with future field developments. '
    'The system trains on historical data to learn which types of proposals proved influential. '
    'This predictive approach complements taste modeling: while taste modeling asks "what would '
    'scientist X find promising?", proposal prediction asks "what will the field find important?"'
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. LLM PERSONALIZATION AND COGNITIVE SIMULATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. LLM Personalization and Cognitive Simulation', level=1)

doc.add_heading('4.1 LaMP: Language Model Personalization', level=2)
doc.add_paragraph(
    'The LaMP benchmark [10] (Salemi et al., 2023) provides a systematic framework for evaluating '
    'personalized text generation across seven tasks spanning classification and generation. LaMP '
    'uses retrieval-augmented approaches where a user\'s historical outputs serve as context. '
    'Evaluation employs accuracy, F1, ROUGE, and text similarity metrics. LaMP\'s methodology is '
    'directly applicable to taste modeling: a scientist\'s published papers, reviews, and '
    'correspondence serve as the "user profile" from which taste preferences can be extracted.'
)

doc.add_heading('4.2 Personalization Methods Taxonomy', level=2)
doc.add_paragraph(
    'Chen et al. [31] (2024) provide a comprehensive taxonomy of LLM personalization methods: '
    '(1) profile-based prompting\u2014injecting user profiles and style examples into prompts; '
    '(2) retrieval-augmented personalization\u2014selecting relevant user history items to condition '
    'generation (as in LaMP); (3) fine-tuning with personal data\u2014using LoRA and parameter-efficient '
    'methods on user-specific corpora; (4) preference learning via RLHF/DPO\u2014training on preference '
    'pairs to align model outputs with individual taste. Direct Preference Optimization (DPO) [32] '
    '(Rafailov et al., 2023, NeurIPS) is particularly relevant: it trains models on preference pairs '
    'without a reward model, and can be extended to personalized settings where different users '
    'have different preferences.'
)

doc.add_heading('4.3 Cognitive Simulation Beyond Behavioral Imitation', level=2)
doc.add_paragraph(
    '"Can Large Language Models Simulate Human Cognition Beyond Behavioral Imitation?" [12] (Shu '
    'et al., 2024) critically examines whether LLMs can genuinely model human cognitive processes '
    'or merely reproduce surface-level behavioral patterns. The authors identify key challenges: '
    '(a) the underspecification problem\u2014many different cognitive processes can produce the same '
    'behavioral output; (b) the training data gap\u2014LLMs are trained on text outputs, not cognitive '
    'processes; (c) the evaluation problem\u2014how to verify genuine cognitive modeling vs. surface '
    'correlations. These challenges directly apply to scientific taste modeling and motivate our '
    'approach of grounding taste models in historical evidence rather than pure behavioral fitting.'
)

doc.add_heading('4.4 MirrorMind and Persona Modeling', level=2)
doc.add_paragraph(
    'MirrorMind [13] (2024) creates LLM-based "cognitive mirrors" reflecting individual users\' '
    'thinking patterns, personality traits, and decision-making styles by combining structured '
    'self-report data with behavioral observation. Park et al.\'s Generative Agents [14] (UIST 2023) '
    'established the foundational paradigm of LLM-based persona simulation with memory architectures. '
    'AgentReview [33] (2024) simulates peer review using LLM agents with different reviewer personas. '
    'While these systems focus on general cognitive traits rather than scientific taste specifically, '
    'their methodology provides a useful template.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. COMPUTATIONAL AESTHETICS AND PREFERENCE LEARNING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Computational Aesthetics and Preference Learning', level=1)

doc.add_heading('5.1 Neuroaesthetics and the Aesthetic Triad', level=2)
doc.add_paragraph(
    'Neuroaesthetics research [34] (Ramachandran & Hirstein, 1999; Vessel et al., 2012, 2018) '
    'has identified neural correlates of aesthetic experience. The "aesthetic triad" model '
    'decomposes aesthetic judgment into three interacting systems: sensory-motor processing, '
    'knowledge-meaning integration, and emotion-valuation. This framework suggests that scientific '
    'taste similarly operates through the integration of domain knowledge (understanding a theory\'s '
    'technical content), aesthetic processing (perceiving mathematical elegance), and evaluative '
    'judgment (assessing significance and promise).'
)

doc.add_heading('5.2 Computational Models of Aesthetic Judgment', level=2)
doc.add_paragraph(
    'Berlyne\'s (1971) "Wundt Curve" [35] models preference as an inverted-U function of stimulus '
    'complexity/arousal, suggesting that scientists prefer theories of intermediate complexity\u2014'
    'neither trivially simple nor incomprehensibly complex. Machado et al. [36] (2008) formalized '
    'aesthetic properties (complexity, order, novelty) mathematically. Bayesian models of aesthetic '
    'preference [37] (extending Tenenbaum and colleagues\' work) model preferences as posterior '
    'beliefs combining prior taste with observed features\u2014a framework directly applicable to '
    'modeling how a scientist\'s taste interacts with the features of a candidate theory.'
)

doc.add_heading('5.3 Preference Learning', level=2)
doc.add_paragraph(
    'The field of preference learning [38] (F\u00fcrnkranz & H\u00fcllermeier, 2010) provides '
    'computational tools for learning preferences from pairwise comparisons, rankings, and ratings. '
    'These methods are directly applicable to taste modeling: given documented cases where a '
    'scientist chose between alternatives, preference learning algorithms can infer the underlying '
    'evaluation function. The pairwise ranking loss used in our proposed taste calibration (Section '
    '6.5) is a standard preference learning technique.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. GAPS, SYNTHESIS, AND PROPOSED FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Critical Gaps, Synthesis, and Proposed Framework', level=1)

doc.add_heading('6.1 Identified Gaps', level=2)

p = doc.add_paragraph()
run = p.add_run('Gap 1: No Formal Model of Individual Scientific Taste. ')
run.bold = True
p.add_run(
    'Despite extensive philosophical discussion and growing computational capabilities, no system '
    'models "research taste" as a personal attribute of individual scientists. Existing systems '
    '(AI Scientist, MooseChem, VirtualLab) create generic scientist roles, not personalized models.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 2: Personalization Methods Not Applied to Scientific Judgment. ')
run.bold = True
p.add_run(
    'Personalization methods are well-developed (LaMP, RAG, DPO) but have not been applied to '
    'scientific evaluation. Bridging personalized LLMs with scientific judgment is a novel direction.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 3: No Benchmark for Scientific Taste Fidelity. ')
run.bold = True
p.add_run(
    'Existing benchmarks evaluate idea generation (Si et al.) or personalization (LaMP) separately, '
    'but none evaluate whether a model accurately captures a specific scientist\'s preferences.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 4: Evidence Grounding vs. Hallucination. ')
run.bold = True
p.add_run(
    'LLM-based persona simulation relies on behavioral fitting without requiring historical '
    'evidence to support claims. For historical scientists, evidence grounding is essential.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 5: Temporal Dynamics. ')
run.bold = True
p.add_run(
    'Scientists\' taste evolves (e.g., Einstein\'s methodological shift around 1920 [3]). No '
    'existing system models this temporal evolution or enforces anachronism prevention.'
)

p = doc.add_paragraph()
run = p.add_run('Gap 6: Cognitive Science Frameworks Unapplied. ')
run.bold = True
p.add_run(
    'Computational aesthetics provides theoretical frameworks (Bayesian preference models, aesthetic '
    'triad, preference learning from pairwise comparisons) that could inform taste modeling but '
    'have not been applied to scientific decision-making.'
)

doc.add_heading('6.2 Proposed Framework', level=2)
doc.add_paragraph(
    'We propose a framework for computational research taste modeling that addresses these gaps:'
)

items = [
    '(1) Taste Axes: Following Holton\'s thematic analysis, define taste as a vector in a '
    'multi-dimensional space of scientific preferences, with each axis grounded in evidence.',
    '(2) Evidence-Based Scoring: Derive axis weights from historical evidence rather than '
    'behavioral fitting, using RAG to ground each evaluation in specific sources.',
    '(3) Temporal Dynamics: Support cutoffs and period-specific weight adjustments to capture '
    'how a scientist\'s taste evolves over their career.',
    '(4) Preference Learning Calibration: Refine weights using documented choice data via '
    'pairwise ranking loss, following the preference learning paradigm [38].',
    '(5) Explicit Inference Marking: Strictly distinguish evidence-based scores from model '
    'inferences, addressing the transparency gap.',
    '(6) Taste Benchmark: Construct a benchmark of documented preferences with difficulty '
    'stratification and temporal sensitivity testing.',
]
for item in items:
    doc.add_paragraph(item)

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    'Scientific research taste represents a critical yet computationally underexplored dimension '
    'of scientific cognition. While philosophers have long recognized its importance (Polanyi, '
    'Holton, Kuhn) and AI researchers have built increasingly capable scientific discovery systems '
    '(AI Scientist, SciMON, MooseChem, EvoScientist), the specific challenge of modeling individual '
    'scientists\' taste profiles remains open. The convergence of LLM capabilities, personalization '
    'methods (LaMP, DPO), preference learning frameworks, and rich historical scholarship on eminent '
    'scientists creates a unique opportunity. We believe that evidence-based, multi-dimensional taste '
    'modeling\u2014grounded in historical sources, calibrated through preference learning, and transparent '
    'about its limitations\u2014represents a promising direction for both AI-assisted scientific '
    'discovery and history/philosophy of science. By operationalizing Holton\'s themata as '
    'computationally measurable taste axes, and leveraging the RAG paradigm to ground evaluations '
    'in specific historical evidence, we can begin to bridge the gap between philosophical insight '
    'and computational capability.'
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading('References', level=1)

refs = [
    '[1] Feynman, R. P. (1974). "Cargo Cult Science." Caltech Commencement Address.',
    '[2] Pais, A. (1982). Subtle is the Lord: The Science and the Life of Albert Einstein. Oxford University Press.',
    '[3] van Dongen, J. (2010). Einstein\'s Unification. Cambridge University Press.',
    '[4] Medawar, P. B. (1967). The Art of the Soluble. Methuen & Co.',
    '[5] Polanyi, M. (1958). Personal Knowledge: Towards a Post-Critical Philosophy. University of Chicago Press.',
    '[6] Kuhn, T. S. (1962). The Structure of Scientific Revolutions. University of Chicago Press.',
    '[7] Holton, G. (1973). Thematic Origins of Scientific Thought: Kepler to Einstein. Harvard University Press.',
    '[8] Lu, C., Lu, C., Lange, R. T., Foerster, J., Clune, J., & Ha, D. (2024). "The AI Scientist: Towards Fully Automated Open-Ended Scientific Discovery." arXiv:2408.06292.',
    '[9] Radensky, M., et al. (2023). "SciMON: Scientific Inspiration Machines Optimized for Novelty." ACL 2023.',
    '[10] Salemi, A., Mysore, S., Bendersky, M., & Zamani, H. (2023). "LaMP: When Large Language Models Meet Personalization." arXiv:2304.11406.',
    '[11] Si, C., Yang, D., & Hashimoto, T. (2024). "Can LLMs Generate Novel Research Ideas? A Large-Scale Human Study with 100+ NLP Researchers." arXiv:2409.04109.',
    '[12] Shu, T., et al. (2024). "Can Large Language Models Simulate Human Cognition Beyond Behavioral Imitation?" arXiv.',
    '[13] MirrorMind. (2024). LLM-based Cognitive Mirror for Personality and Cognitive Modeling. arXiv.',
    '[14] Park, J. S., O\'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P., & Bernstein, M. S. (2023). "Generative Agents: Interactive Simulacra of Human Behavior." UIST 2023.',
    '[15] Swanson, K., et al. (2024). "VirtualLab: AI Agents as Researchers in Nanobody Engineering." arXiv.',
    '[16] Kant, I. (1790). Critique of Judgment. Trans. Werner Pluhar, Hackett Publishing.',
    '[17] Bourdieu, P. (1984). Distinction: A Social Critique of the Judgement of Taste. Harvard University Press.',
    '[18] Polanyi, M. (1966). The Tacit Dimension. University of Chicago Press.',
    '[19] Holton, G. (1998). The Scientific Imagination. Harvard University Press.',
    '[20] "AI Can Learn Scientific Taste." (2024). Perspective article.',
    '[21] Li, Z., et al. (2024). "MooseChem: Large Language Models for Rediscovering Unseen Chemistry Scientific Hypotheses." arXiv.',
    '[22] Radensky, M., et al. (2024). "Scideator: Human-LLM Scientific Idea Generation Grounded in Research-Paper Facet Recombination." arXiv.',
    '[23] Zhang, C., et al. (2024). "EvoScientist: Evolving Scientific Hypotheses with Language Models." arXiv.',
    '[24] SciMuse. (2024). Cross-Domain Scientific Inspiration System.',
    '[25] Baek, J., et al. (2024). "ResearchAgent: Iterative Research Idea Generation over Scientific Literature with Large Language Models." arXiv:2404.07738.',
    '[26] Ghafarollahi, A. & Buehler, M. J. (2024). "SciAgents: Automating Scientific Discovery through Multi-Agent Intelligent Graph Reasoning." arXiv.',
    '[27] Sprueill, H., et al. (2024). "ChemReasoner: Heuristic Search over a Large Language Model\'s Knowledge Space for Catalyst Discovery." ICML 2024.',
    '[28] Swanson, K., et al. (2024). "VirtualLab: AI Agents as Researchers in Nanobody Engineering." arXiv.',
    '[29] Liang, W., et al. (2024). "Can Large Language Models Provide Useful Feedback on Research Papers? A Large-Scale Empirical Analysis." arXiv:2310.01783.',
    '[30] Ye, J., et al. (2024). "Learning to Predict Future-Aligned Research Proposals with Language Models." arXiv.',
    '[31] Chen, J., et al. (2024). "A Survey on Personalized Large Language Models." arXiv.',
    '[32] Rafailov, R., Sharma, A., Mitchell, E., Ermon, S., Manning, C. D., & Finn, C. (2023). "Direct Preference Optimization: Your Language Model is Secretly a Reward Model." NeurIPS 2023.',
    '[33] AgentReview. (2024). "Exploring Peer Review Dynamics with LLM Agents." arXiv.',
    '[34] Ramachandran, V. S. & Hirstein, W. (1999). "The Science of Art: A Neurological Theory of Aesthetic Experience." Journal of Consciousness Studies, 6(6-7).',
    '[35] Berlyne, D. E. (1971). Aesthetics and Psychobiology. Appleton-Century-Crofts.',
    '[36] Machado, P., et al. (2008). "Computational Aesthetics." In Computational Intelligence, Vol. 1.',
    '[37] Tenenbaum, J. B., Kemp, C., Griffiths, T. L., & Goodman, N. D. (2011). "How to Grow a Mind: Statistics, Structure, and Abstraction." Science, 331(6022).',
    '[38] F\u00fcrnkranz, J. & H\u00fcllermeier, E. (2010). Preference Learning. Springer.',
    '[39] Einstein, A. (1933). "On the Method of Theoretical Physics." Herbert Spencer Lecture, Oxford.',
    '[40] Einstein, A. (1949). "Autobiographical Notes." In P.A. Schilpp (ed.), Albert Einstein: Philosopher-Scientist.',
    '[41] Fine, A. (1986). The Shaky Game: Einstein, Realism, and the Quantum Theory. University of Chicago Press.',
    '[42] Norton, J. (1984). "How Einstein Found His Field Equations." Historical Studies in the Physical Sciences.',
    '[43] Howard, D. (1985). "Einstein on Locality and Separability." Studies in History and Philosophy of Science.',
    '[44] Isaacson, W. (2007). Einstein: His Life and Universe. Simon & Schuster.',
    '[45] Kang, D., et al. (2018). "A Dataset of Peer Reviews (PeerRead): Collection, Insights and NLP Applications." NAACL 2018.',
    '[46] Wang, H., et al. (2023). "Scientific Discovery in the Age of Artificial Intelligence." Nature, 620.',
    '[47] Huang, Q., et al. (2024). "MLAgentBench: Evaluating Language Agents on Machine Learning Experimentation." ICLR 2024.',
    '[48] Chatterjee, A. (2013). The Aesthetic Brain: How We Evolved to Desire Beauty and Enjoy Art. Oxford University Press.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ────────────────────────────────────────────────────────────────────
output_dir = r"D:\student\scientist_taste\EinsteinResearchTaste\docs"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "literature_review.docx")
doc.save(output_path)
print(f"Updated literature review saved to {output_path}")
print(f"Total references: {len(refs)}")
